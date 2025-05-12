import docx
import sys

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    
    print(f"Document has {len(doc.paragraphs)} paragraphs")
    print(f"Document has {len(doc.tables)} tables")
    
    # Extract text from paragraphs
    print("\n=== DOCUMENT CONTENT ===\n")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"Paragraph {i}: {para.text}")
    
    # Extract text from tables
    if doc.tables:
        print("\n=== TABLES ===\n")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i}:")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                print(" | ".join(row_text))

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <path_to_docx_file>")
        sys.exit(1)
    
    extract_text_from_docx(sys.argv[1])
