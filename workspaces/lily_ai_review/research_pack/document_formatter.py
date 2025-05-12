import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from .table_generator import TableGenerator
import re
from app.membership.config import get_tier_config
from app.services.paper_generator.config import get_section_word_count
from typing import Dict
import textwrap
import io
import string
import random
import unicodedata

logger = logging.getLogger(__name__)

class DocumentFormatter:
    def __init__(self):
        self.document = Document()
        self._setup_document_styles(self.document)
        self._setup_sections()
        self.table_generator = TableGenerator()
        # Add section counters for proper numbering
        self.section_counters = {1: 0, 2: 0, 3: 0}

    def _setup_document_styles(self, doc: Document):
        """Setup document styles for a professional academic paper"""
        # Normal text style
        if 'Normal' in doc.styles:
            normal_style = doc.styles['Normal']
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.space_after = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5

        # Use built-in heading styles
        if 'Heading 1' in doc.styles:
            heading1 = doc.styles['Heading 1']
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.paragraph_format.space_before = Pt(24)
            heading1.paragraph_format.space_after = Pt(12)

        if 'Heading 2' in doc.styles:
            heading2 = doc.styles['Heading 2']
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.paragraph_format.space_before = Pt(18)
            heading2.paragraph_format.space_after = Pt(6)

        if 'Heading 3' in doc.styles:
            heading3 = doc.styles['Heading 3']
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.paragraph_format.space_before = Pt(12)
            heading3.paragraph_format.space_after = Pt(6)

        # Title style
        if 'Title' in doc.styles:
            title_style = doc.styles['Title']
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)

        # TOC styles
        if 'TOC Heading' not in doc.styles:
            toc_style = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(16)
            toc_style.font.bold = True
            toc_style.paragraph_format.space_before = Pt(24)
            toc_style.paragraph_format.space_after = Pt(12)

        if 'TOC' not in doc.styles:
            toc_style = doc.styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(12)
            toc_style.paragraph_format.left_indent = Inches(0.25)
            toc_style.paragraph_format.space_after = Pt(6)

        if 'TOC 1' not in doc.styles:
            toc1_style = doc.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
            toc1_style.font.size = Pt(12)
            toc1_style.paragraph_format.left_indent = Inches(0.25)
            toc1_style.paragraph_format.space_after = Pt(6)

        if 'TOC 2' not in doc.styles:
            toc2_style = doc.styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
            toc2_style.font.size = Pt(11)
            toc2_style.paragraph_format.left_indent = Inches(0.5)
            toc2_style.paragraph_format.space_after = Pt(6)

    def _setup_sections(self):
        """Setup sections for headers and footers"""
        section = self.document.sections[0]

        # Add page numbers and Lily encouragement to footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add random Lily encouragement statement
        encouragements = [
            "Lily is here to help you succeed in your academic journey! 💫",
            "Remember, every great researcher started with a single question. Keep going! 💫",
            "You're making real progress—don't be afraid to explore new ideas! 💫",
            "Every challenge is a chance to learn. Stay curious! 💫",
            "You've got this! Your unique perspective matters. 💫",
            "Keep pushing forward—breakthroughs come from persistence! 💫",
            "Celebrate your progress—each step brings you closer to your goal! 💫"
        ]
        encouragement = random.choice(encouragements)
        footer_para.text = encouragement + "   |   Page "

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        # Style the footer text
        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        # Create a table of contents style
        if 'TOC Heading' not in self.document.styles:
            toc_style = self.document.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(16)
            toc_style.font.bold = True
            toc_style.paragraph_format.space_before = Pt(24)
            toc_style.paragraph_format.space_after = Pt(12)

        if 'TOC 1' not in self.document.styles:
            toc1_style = self.document.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
            toc1_style.font.size = Pt(12)
            toc1_style.paragraph_format.left_indent = Inches(0.25)
            toc1_style.paragraph_format.space_after = Pt(6)

        if 'TOC 2' not in self.document.styles:
            toc2_style = self.document.styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
            toc2_style.font.size = Pt(11)
            toc2_style.paragraph_format.left_indent = Inches(0.5)
            toc2_style.paragraph_format.space_after = Pt(6)

    def create_section_divider(self, section_number: int, section_title: str, section_description: str, subsections: list, education_level: str = "university"):
        """Create a section divider page."""
        # Section colors
        section_colors = {
            1: "3498db",  # Blue
            2: "2ecc71",  # Green
            3: "e67e22",  # Orange
            4: "9b59b6",  # Purple
            5: "f1c40f",  # Yellow
        }

        # Add page break
        self.document.add_page_break()

        # Add section number and title
        section_heading = self.document.add_paragraph()
        section_heading.style = 'Heading 1'
        section_run = section_heading.add_run(f"SECTION {section_number}: {section_title}")
        section_run.font.color.rgb = RGBColor.from_string(section_colors.get(section_number, "000000"))
        section_run.font.size = Pt(24)
        section_run.font.bold = True

        # Add horizontal line
        self.document.add_paragraph("_______________________________________________")

        # Add section description
        description_para = self.document.add_paragraph()
        description_para.add_run(section_description).italic = True

        # Add subsections list
        self.document.add_paragraph("In this section:").bold = True
        for subsection in subsections:
            subsection_para = self.document.add_paragraph()
            subsection_para.style = 'List Bullet'
            subsection_para.add_run(subsection).bold = True

        # Add horizontal line
        self.document.add_paragraph("_______________________________________________")

        # Add section footer
        footer_para = self.document.add_paragraph()
        footer_para.add_run(f"Section {section_number} of 5 | {education_level.title()} Research Pack").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create_title_page(self, title: str, author: str, institution: str, date: str, education_level: str = None):
        """Create a professional title page with Lily logo"""
        # Add Lily logo at the top
        try:
            logo_path = "app/static/lilylogo.png"
            if os.path.exists(logo_path):
                logo_para = self.document.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2))
                logger.info(f"Added Lily logo from {logo_path}")
            else:
                logger.warning(f"Lily logo not found at {logo_path}")
        except Exception as e:
            logger.error(f"Error adding Lily logo: {str(e)}")

        # Title - use the actual title, not the topic
        title_para = self.document.add_paragraph(title)
        title_para.style = 'Title'

        # Add some spacing
        self.document.add_paragraph()

        # Author information - use the actual name, not "Researcher"
        author_para = self.document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"Prepared for\n{author}").bold = True

        # Institution
        inst_para = self.document.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_para.add_run(institution)

        # Date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(date)

        # Education Level (if provided)
        if education_level:
            # Format the education level for display (e.g., "high-school" -> "High School Level")
            formatted_level = education_level.replace('-', ' ').title() + " Level"

            # Add education level with distinctive formatting
            level_para = self.document.add_paragraph()
            level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            level_run = level_para.add_run(formatted_level)
            level_run.bold = True
            level_run.font.size = Pt(14)

            # Add a colored box around the education level
            level_para.paragraph_format.border_between = True
            level_para.paragraph_format.border_width = Pt(1)

            # Add a note about the education level
            note_para = self.document.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_para.add_run("This document has been adapted to the specified education level.").italic = True

        # Add page break after title page
        self.document.add_page_break()

    def add_table_of_contents(self):
        """Add a table of contents that will update correctly when the document is opened"""
        # Add a heading for the TOC
        toc_heading = self.document.add_heading("Table of Contents", level=1)
        toc_heading.style = 'Heading 1'
        # Don't number the TOC
        toc_heading._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 0

        # Add a paragraph for the TOC field
        paragraph = self.document.add_paragraph()

        # Create the TOC field - this needs to be properly structured to work
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Begin field
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)

        # Field instruction text
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        # This specific format of the TOC field code is important
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._element.append(instrText)

        # Separate field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar2)

        # Add placeholder text
        placeholder_run = paragraph.add_run("Right-click to update field.")

        # End field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar3)

        # Add a page break after TOC
        self.document.add_page_break()

    def clean_markdown(self, text: str) -> str:
        """Clean markdown syntax from text"""
        if not text:
            return ""

        # Remove heading markers more aggressively
        text = text.strip()
        while '#' in text:
            text = text.replace('#', '')

        # Clean other markdown syntax
        replacements = {
            '**': '',
            '__': '',
            '*': '',
            '_': '',
            '```': '',
            '`': '',
            'Figure:': '',
            'figure:': '',
            '[': '',
            ']': '',
            '  ': ' '  # Replace double spaces with single
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        # Clean up multiple spaces and lines
        text = ' '.join(text.split())
        return text.strip()

    def add_section(self, heading: str, content: str, level: int = 1):
        """Add a new section with heading and content"""
        # Add page break before each main section (level 1)
        if level == 1 and self.document.paragraphs:
            self.document.add_page_break()

        # Clean up markdown-style headings
        heading = self.clean_markdown(heading)
        heading = heading.lstrip('#').strip()

        # Handle section numbering
        if level <= 3:  # We only number up to level 3
            if level == 1:
                # Increment level 1 counter, reset others
                self.section_counters[1] += 1
                self.section_counters[2] = 0
                self.section_counters[3] = 0
                section_number = f"{self.section_counters[1]}."
            elif level == 2:
                # Increment level 2 counter, reset level 3
                self.section_counters[2] += 1
                self.section_counters[3] = 0
                section_number = f"{self.section_counters[1]}.{self.section_counters[2]}"
            elif level == 3:
                # Increment level 3 counter
                self.section_counters[3] += 1
                section_number = f"{self.section_counters[1]}.{self.section_counters[2]}.{self.section_counters[3]}"

            # Add numbered heading (except for special sections)
            special_sections = ["references", "abstract", "key points", "appendix", "acknowledgements", "table of contents"]
            if heading.lower() not in special_sections:
                heading_with_number = f"{section_number} {heading}"
            else:
                heading_with_number = heading
                if level == 1:
                    self.section_counters[1] -= 1
        else:
            heading_with_number = heading

        # Add extra spacing before section heading
        if level == 1:
            self.document.add_paragraph().paragraph_format.space_before = Pt(24)

        # Add the section heading with appropriate style
        section_heading = self.document.add_heading(heading_with_number, level=level)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply the proper heading style
        if level == 1:
            section_heading.style = 'Heading 1'
        elif level == 2:
            section_heading.style = 'Heading 2'
        elif level == 3:
            section_heading.style = 'Heading 3'
        else:
            section_heading.style = 'Heading 1'

        # Process content and add paragraphs with proper formatting
        if content:
            # Split content into paragraphs
            paragraphs = content.split('\n\n')

            # Add each paragraph with proper spacing
            for i, para_text in enumerate(paragraphs):
                if para_text.strip():
                    # Clean the paragraph text
                    para_text = self.clean_markdown(para_text)

                    # Add the paragraph
                    paragraph = self.document.add_paragraph()
                    paragraph.text = para_text.strip()

                    # Set paragraph formatting
                    paragraph.style = 'Normal'
                    paragraph.paragraph_format.first_line_indent = Inches(0.5)
                    paragraph.paragraph_format.space_after = Pt(12)

                    # Add extra spacing between subsections if detected
                    if i > 0 and para_text.strip().startswith(str(self.section_counters[1])):
                        paragraph.paragraph_format.space_before = Pt(18)

        # Add extra spacing after the section
        if level == 1:
            self.document.add_paragraph().paragraph_format.space_after = Pt(24)

    def add_bullet_point(self, text: str):
        """Add a bullet point to the document with proper formatting"""
        paragraph = self.document.add_paragraph(style='List Bullet')
        paragraph.text = text.strip()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        return paragraph

    def add_heading(self, text: str, level: int = 1):
        """Add a heading with the specified level"""
        if not text:
            return None

        # Clean markdown formatting
        text = self.clean_markdown(text)

        # Create the heading
        heading = self.document.add_heading(text, level=level)

        # Apply appropriate formatting based on level
        if level == 0:  # Major section divider
            heading.style = self.document.styles['Title']
            heading.paragraph_format.space_before = Pt(24)
            heading.paragraph_format.space_after = Pt(12)
        elif level == 1:
            heading.style = self.document.styles['Heading 1']
        elif level == 2:
            heading.style = self.document.styles['Heading 2']
        elif level == 3:
            heading.style = self.document.styles['Heading 3']

        return heading

    def add_paragraph(self, text: str = ""):
        """Add a paragraph with the specified text, processing any embedded Lily insights"""
        if not text:
            return self.document.add_paragraph()

        # Clean markdown formatting
        text = self.clean_markdown(text)

        # Check for Lily's embedded notes
        lily_note_pattern = r'\[Lily\'s (note|suggestion|question|insight|tip): (.*?)\]'
        lily_notes = re.findall(lily_note_pattern, text)

        if not lily_notes:
            # No Lily notes, add paragraph as normal
            paragraph = self.document.add_paragraph()
            paragraph.text = text
            paragraph.style = 'Normal'
            return paragraph

        # Process text with Lily notes
        parts = re.split(lily_note_pattern, text)

        # The first part is text before any Lily note
        if parts[0]:
            paragraph = self.document.add_paragraph()
            paragraph.text = parts[0].strip()
            paragraph.style = 'Normal'

        # Process the rest in triplets: note_type, note_content, text_after
        i = 1
        while i < len(parts) - 1:
            note_type = parts[i]
            note_content = parts[i+1]

            # Map note type to insight type
            insight_type = "standard"
            if "citation" in note_type.lower() or "reference" in note_type.lower():
                insight_type = "citation"
            elif "idea" in note_type.lower() or "suggestion" in note_type.lower():
                insight_type = "idea"
            elif "counter" in note_type.lower() or "alternative" in note_type.lower():
                insight_type = "counterpoint"
            elif "confidence" in note_type.lower():
                insight_type = "confidence"
            elif "warning" in note_type.lower() or "caution" in note_type.lower():
                insight_type = "warning"
            elif "tip" in note_type.lower() or "writing" in note_type.lower():
                insight_type = "tip"
            elif "question" in note_type.lower() or "consider" in note_type.lower():
                insight_type = "question"

            # Add Lily's insight
            self.add_lily_insight(note_content, insight_type)

            # Add following text if present
            after_text = parts[i+2] if i+2 < len(parts) else ""
            if after_text.strip():
                after_para = self.document.add_paragraph()
                after_para.text = after_text.strip()
                after_para.style = 'Normal'

            i += 3

        # Return the last paragraph added
        return self.document.paragraphs[-1]

    def add_page_break(self):
        """Add a page break to the document"""
        self.document.add_page_break()

    def add_figure(self, image_path: str, caption: str):
        """Add a figure with caption and maintain aspect ratio"""
        try:
            if not os.path.exists(image_path):
                logger.error(f"Image not found at path: {image_path}")
                return

            # Add spacing before figure
            self.document.add_paragraph().paragraph_format.space_before = Pt(12)

            # Center the figure
            figure_para = self.document.add_paragraph()
            figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure_run = figure_para.add_run()
            picture = figure_run.add_picture(image_path)

            # Set width while maintaining aspect ratio
            target_width = Inches(6.0)  # Standard width for figures
            aspect_ratio = picture.height / picture.width
            picture.width = target_width
            picture.height = int(target_width * aspect_ratio)

            # Add figure number
            self.figure_count = getattr(self, 'figure_count', 0) + 1

            # Create or get the Caption style
            if 'Caption' not in self.document.styles:
                caption_style = self.document.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                caption_style.font.size = Pt(10)
                caption_style.font.italic = False
                caption_style.font.bold = True
                caption_style.paragraph_format.space_before = Pt(6)
                caption_style.paragraph_format.space_after = Pt(18)
                caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Clean up caption and add with automatic numbering
            clean_caption = caption.replace('Figure:', '').strip()  # Remove any existing 'Figure:' prefix
            clean_caption = clean_caption.replace('##', '').strip()  # Remove markdown heading syntax
            caption_para = self.document.add_paragraph()
            caption_para.style = 'Caption'
            caption_para.add_run(f'Figure {self.figure_count}. {self.clean_markdown(caption)}')

            # Add spacing after figure
            self.document.add_paragraph().paragraph_format.space_after = Pt(12)

            logger.info(f"Successfully added figure from {image_path}")
        except Exception as e:
            logger.error(f"Error adding figure: {str(e)}")
            logger.error(f"Image path: {image_path}")

    def add_header_footer(self, title: str):
        """Add header and footer with page numbers to the document"""
        # Get the section
        section = self.document.sections[0]

        # Add title to header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = title
        header_para.style = self.document.styles['Normal']
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page numbers to footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save_document(self, filepath: str) -> str:
        """Save the document to the specified path"""
        try:
            logger.info(f"Attempting to save document to {filepath}")

            # Sanitize the filepath to remove URL-unsafe characters
            original_filepath = filepath
            directory = os.path.dirname(filepath)
            filename = os.path.basename(filepath)

            # Replace URL-unsafe characters with underscores
            sanitized_filename = self.sanitize_filename(filename)

            # Construct the sanitized filepath
            sanitized_filepath = os.path.join(directory, sanitized_filename)

            if sanitized_filepath != original_filepath:
                logger.info(f"Sanitized filepath from {original_filepath} to {sanitized_filepath}")

            # Handle case when no directory is specified (save in current directory)
            if directory:
                logger.info(f"Creating directory: {directory}")
                os.makedirs(directory, exist_ok=True)
                logger.info("Directory created successfully")

            # Save the document
            logger.info("Saving document...")
            self.document.save(sanitized_filepath)
            logger.info(f"Document saved successfully to {sanitized_filepath}")
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return False

    def sanitize_filename(self, filename: str) -> str:
        """Sanitize a filename to make it safe for saving"""
        # Replace URL-unsafe characters with underscores
        sanitized = re.sub(r'[?&+=#%:;/\\,\'\"\-]', '_', filename)
        # Replace multiple underscores with a single one
        sanitized = re.sub(r'_{2,}', '_', sanitized)
        # Limit the length of the filename
        if len(sanitized) > 100:
            base, ext = os.path.splitext(sanitized)
            sanitized = base[:96] + ext
        return sanitized

    def generate_template(self, topic: str, section: str, key_points: list, tier_name: str = None):
        """
        Generate a template prompt for section generation with tier-appropriate word counts

        Args:
            topic: The topic of the paper
            section: The section to generate
            key_points: List of key points to address
            tier_name: The subscription tier name (premium, standard, or sample)
        """
        key_points_str = "\n".join([f"- {point}" for point in key_points])

        # Get appropriate word count for this section and tier
        section_word_count = get_section_word_count(section, tier_name)
        word_count_instruction = f"Target word count for this section: {section_word_count} words."

        # Add tier-specific depth instructions
        tier_specific_guidance = ""
        if tier_name:
            if tier_name.lower() == "premium":
                tier_specific_guidance = """
                Premium Tier Requirements:
                - Provide comprehensive, in-depth analysis with substantial detail
                - Include multiple perspectives and nuanced arguments
                - Use extensive evidence and examples to support claims
                - Address potential counterarguments or limitations
                - Ensure thorough coverage of all aspects of the topic
                """
            elif tier_name.lower() == "sample":
                tier_specific_guidance = """
                Sample Tier Requirements:
                - Focus on the most essential points only
                - Be concise and direct in your writing
                - Provide a representative example of quality without depth
                - Maintain academic tone while being brief
                """

        prompt = f"""You are a PhD-level researcher and expert in {topic}. Write a rigorous, publication-quality {section} section for an academic paper. Do not use markdown formatting or special characters in headings.

{word_count_instruction}

Key points to address:
{key_points_str}

{tier_specific_guidance}

Guidelines:
1. Content Requirements:
   - Provide comprehensive coverage of the topic
   - Present clear, evidence-based arguments
   - Include relevant statistics and data when applicable
   - Address current developments in the field

2. Academic Standards:
   - Use formal academic language
   - Support claims with peer-reviewed sources
   - Follow logical flow of ideas
   - Maintain objective tone throughout

3. Citations and References:
   - Include in-text citations [Author, Year]
   - Reference recent and seminal works
   - Cite primary research studies
   - Use authoritative sources

4. Section-Specific Focus:
   For {section}:
   - Follow standard conventions for this section type
   - Provide appropriate depth and detail
   - Connect to the paper's overall narrative
   - Address key controversies or debates if relevant

5. Tables and Figures:
   - IMPORTANT: When presenting data or comparisons, use well-formatted tables with real data
   - Format tables using pipe notation (|) for clear column separation
   - Include a header row and a separator row below it
   - Example table format:

   | Variable | Value | Significance |
   |----------|-------|-------------|
   | Item 1   | 85%   | p < 0.01    |
   | Item 2   | 72%   | p < 0.05    |

Write the {section} section. Use clear headings without markdown symbols and make sure to include numbered sections and subsections. When appropriate, include properly formatted tables with relevant data."""
        return prompt

    def format_research_pack(self, research_pack: Dict) -> Document:
        """
        Format a complete research pack into a single well-organized document.
        """
        doc = Document()

        # Set up document styles
        self._setup_document_styles(doc)

        # Create a style for Lily's guidance text
        if 'Lily Guide' not in doc.styles:
            guide_style = doc.styles.add_style('Lily Guide', WD_STYLE_TYPE.PARAGRAPH)
            guide_style.font.size = Pt(11)
            guide_style.font.italic = True
            guide_style.font.bold = True
            guide_style.paragraph_format.space_after = Pt(12)
            guide_style.paragraph_format.space_before = Pt(6)

        # Create a style for small text
        if 'Small Text' not in doc.styles:
            small_style = doc.styles.add_style('Small Text', WD_STYLE_TYPE.PARAGRAPH)
            small_style.font.size = Pt(8)
            small_style.font.italic = True
            small_style.paragraph_format.space_after = Pt(6)

        # Title Page
        title = research_pack["example_paper"]["title"]
        doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run("Research Pack").bold = True
        doc.add_paragraph(f"Generated: {research_pack['date_generated']}")

        # Add banner text with Lily's style
        banner = doc.add_paragraph(style='Lily Guide')
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner.add_run(f"""Hey there! I'm Lily 👋
Your AI Research Partner

I've analyzed your topic "{research_pack['topic']}" and prepared this comprehensive research pack just for you! Inside you'll find everything you need to create an amazing paper:
• A well-structured example paper to inspire your writing
• A personalized research roadmap tailored to your topic
• Topic-specific study materials to deepen your understanding
• Custom enhancement suggestions to make it uniquely yours

Remember: This is your starting point - I'm here to guide you, but your unique perspective is what will make this paper special! Use this pack to:
✓ Get a solid grasp of the topic's structure
✓ Follow a clear research strategy
✓ Develop your own ideas and arguments
✓ Create something you're proud of!""")

        doc.add_page_break()

        # SECTION 1: RESEARCH PLANNING GUIDE
        doc.add_heading("RESEARCH PLANNING GUIDE", 0)

        # 1. Understanding Your Topic
        doc.add_heading("1. Understanding Your Topic", 1)
        topic_guide = doc.add_paragraph(style='Lily Guide')
        topic_guide.add_run(f"""Let's break down your fascinating topic: "{research_pack['topic']}"

I've identified these key areas for you to focus on:
• Core concepts: {research_pack.get('key_concepts', ['[Key concepts will be listed here]'])[0]}
• Main theories: {research_pack.get('main_theories', ['[Main theories will be listed here]'])[0]}
• Current debates: What experts are discussing right now
• Research gaps: Where you can add your voice

Your Research Goals (I'll help you with each one!):
1. Master the fundamental concepts of {research_pack['topic'].split()[0]}
2. Understand the main arguments in this field
3. Find solid evidence to support your points
4. Develop your unique perspective on the topic""")

        # 2. Research Approach
        doc.add_heading("2. Research Approach", 1)
        approach = doc.add_paragraph(style='Lily Guide')
        approach.add_run(f"""I've created this step-by-step plan specifically for {research_pack['topic']}:

Phase 1: Getting Started (2-3 days)
• Review the example paper I've prepared
• Highlight terms you want to explore further
• Note any questions that come to mind
• List the main arguments you find interesting

Phase 2: Deep Research (4-5 days)
• Search for recent papers on {research_pack['topic']}
• Look for supporting evidence
• Find different viewpoints
• Gather relevant statistics

Phase 3: Critical Analysis (2-3 days)
• Compare what different experts say
• Evaluate how strong the evidence is
• Look for gaps you could address
• Form your own position

Phase 4: Writing Your Paper (4-5 days)
• Create your unique outline
• Write each section
• Add your sources
• Polish your work""")

        # 3. Research Resources
        doc.add_heading("3. Research Resources", 1)
        resources = doc.add_paragraph(style='Lily Guide')
        resources.add_run(f"""I've found these resources particularly helpful for {research_pack['topic']}:

Must-Check Databases:
• Google Scholar - Great for recent papers
• JSTOR - Perfect for theoretical background
• ScienceDirect - Excellent for empirical studies
• Your university library - Don't forget this goldmine!

Types of Sources You'll Need:
• Recent journal articles (especially from {research_pack.get('key_journals', ['relevant journals'])[0]})
• Academic books on {research_pack['topic'].split()[0]}
• Latest conference proceedings
• Field-specific reports
• Real-world case studies

Pro Tips (from your AI buddy 😊):
✓ Use the citations from our example paper as starting points
✓ Focus on papers from the last 5 years
✓ Save everything interesting you find
✓ Take notes while you read - trust me, it helps!""")

        # 4. Writing Strategy
        doc.add_heading("4. Writing Strategy", 1)
        strategy = doc.add_paragraph(style='Lily Guide')
        strategy.add_run(f"""Let's make your paper on {research_pack['topic']} stand out:

Structuring Your Work:
• Follow the outline in our example
• Use clear section headings
• Build your arguments step by step
• Back everything up with evidence

Making Strong Arguments:
• Show different perspectives on {research_pack['topic']}
• Address potential counterarguments
• Use specific examples from your research
• Include relevant data and statistics

Writing Tips:
• Use academic language (but keep it readable!)
• Explain technical terms clearly
• Connect your ideas smoothly
• Keep your style consistent

Watch Out For (I've got your back!):
✗ Don't just repeat what others say
✗ Avoid claims without evidence
✗ Don't ignore opposing viewpoints
✗ Keep the tone professional""")

        # 5. Using This Research Pack
        doc.add_heading("5. Using This Research Pack", 1)
        usage = doc.add_paragraph(style='Lily Guide')
        usage.add_run(f"""Here's how to get the most out of what I've prepared for you:

Example Paper:
• Use it as inspiration, not a template
• Notice how arguments are structured
• See how sources are integrated
• Think about what you'd do differently

Study Guide:
• Start with the basics of {research_pack['topic']}
• Work through the study questions
• Test your understanding
• Note areas where you want to dig deeper

Research Plan:
• Check out the sources I've suggested
• Add new sources as you find them
• Keep track of your own discoveries
• Update the plan as you go

Enhancement Guide:
• Consider all my suggestions
• Add your own insights
• Strengthen any weak points
• Make it uniquely yours

Remember, I'm here to guide you, but your unique perspective is what will make this paper special! Let's create something amazing together! 🌟""")

        doc.add_page_break()

        # Add Lily's responses to user questions if available
        if "lily_responses" in research_pack and research_pack["lily_responses"]:
            doc.add_heading("Lily's Personalized Guidance", 1)

            # Add introduction to this section
            intro = doc.add_paragraph(style='Lily Guide')
            intro.add_run(f"""Based on your specific questions, here's my personalized guidance for your research on {research_pack['topic']}. I've addressed each of your questions with actionable advice and specific suggestions tailored to your needs.""")

            # Add each question and response
            for question, response in research_pack["lily_responses"].items():
                # Add the question as a heading
                doc.add_heading(f"Q: {question}", 2)

                # Add the response
                response_para = doc.add_paragraph()
                response_para.add_run(response)

                # Add a separator
                doc.add_paragraph("---")

            doc.add_page_break()

        # Contents page and rest of document follows...
        # ... rest of the existing code ...

    def add_lily_insight_with_avatar(self, insight_text, insight_type="standard"):
        """
        Add a Lily insight with a character avatar image and more personality

        This creates a more visually distinct callout box with an avatar-like presentation
        for Lily's insights.
        """
        # Create a 2-column table for avatar and content
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Make borders more subtle - using the correct approach for python-docx
        # Instead of trying to manipulate XML directly, use the table style
        table.style = 'Light Grid'  # Use a built-in light table style

        # Add light blue shading to the cell background
        for cell in table._cells:
            shading_element = cell._element.tcPr.xpath('./w:shd')
            if not shading_element:
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6F0FF"/>')
                cell._element.tcPr.append(shading_element)
            else:
                shading_element[0].set(qn('w:fill'), 'E6F0FF')  # Light blue background

        # Set column widths - narrow for avatar, wider for content
        table.columns[0].width = Inches(0.8)  # Avatar column
        table.columns[1].width = self.document.sections[0].page_width - Inches(2.0)  # Content column

        # Avatar cell (first column)
        avatar_cell = table.cell(0, 0)

        # Determine emoji avatar based on insight type
        avatar = "👩‍🎓 "  # Default avatar emoji - scholar
        if insight_type == "citation":
            avatar = "📚 "
        elif insight_type == "idea":
            avatar = "💭 "
        elif insight_type == "counterpoint":
            avatar = "🤔 "
        elif insight_type == "confidence":
            avatar = "🌟 "
        elif insight_type == "warning":
            avatar = "⚠️ "
        elif insight_type == "tip":
            avatar = "✍️ "
        elif insight_type == "question":
            avatar = "❓ "

        # Add avatar emoji with "Lily" text
        avatar_para = avatar_cell.paragraphs[0]
        avatar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        avatar_run = avatar_para.add_run(f"{avatar}\nLily")
        avatar_run.bold = True
        avatar_run.font.size = Pt(18)  # Larger emoji

        # Content cell (second column)
        content_cell = table.cell(0, 1)

        # Add the insight text with a header
        insight_para = content_cell.paragraphs[0]

        # Add a catchy header with personality
        header_text = "Lily's Insight"
        if insight_type == "citation":
            header_text = "Citation Tip from Lily"
        elif insight_type == "idea":
            header_text = "Lily's Brainstorm"
        elif insight_type == "counterpoint":
            header_text = "Consider This..."
        elif insight_type == "confidence":
            header_text = "Lily's Confidence Booster"
        elif insight_type == "warning":
            header_text = "Lily's Caution Flag"
        elif insight_type == "tip":
            header_text = "Lily's Writing Tip"
        elif insight_type == "question":
            header_text = "Question to Consider"

        # Add a colored header
        header_run = insight_para.add_run(f"{header_text}:\n")
        header_run.bold = True
        header_run.font.size = Pt(12)

        # Add the actual insight text
        text_run = insight_para.add_run(insight_text)
        text_run.italic = True
        text_run.font.size = Pt(11)

        # Add spacing paragraph after the table
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_lily_insight(self, insight_text, insight_type="standard"):
        """Add a Lily insight with appropriate styling and emoji"""
        # Use the new avatar version instead for better styling
        return self.add_lily_insight_with_avatar(insight_text, insight_type)

    def add_research_highlight(self, research_type, source, content, year="2023"):
        """
        Add a research highlight box with citation information.

        Args:
            research_type: Type of research (e.g., "Research Methodology")
            source: Source of the research
            content: The highlight content
            year: Publication year
        """
        # Create a table for the highlight with borders
        table = self.document.add_table(rows=2, cols=1)
        table.style = 'Table Grid'

        # Get the header cell
        header_cell = table.cell(0, 0)

        # Add light yellow shading to the header cell background
        shading_element = header_cell._element.tcPr.xpath('./w:shd')
        if not shading_element:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFE0"/>')
            header_cell._element.tcPr.append(shading_element)
        else:
            shading_element[0].set(qn('w:fill'), 'FFFFE0')  # Light yellow background

        # Add header content
        header_para = header_cell.paragraphs[0]
        header_run = header_para.add_run(f"Research Highlight: {research_type}")
        header_run.bold = True
        header_run.font.size = Pt(12)

        # Get the content cell
        content_cell = table.cell(1, 0)

        # Add content
        content_para = content_cell.paragraphs[0]
        content_para.text = content
        content_para.add_run(f"\n\nSource: {source}, {year}")

        # Add spacing after the highlight
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_info_box(self, title, content, box_type="info"):
        """
        Add a visually distinct information box with a title and content.

        Args:
            title: The title of the info box
            content: The main text content
            box_type: Type of box (info, warning, tip, research, fact)
        """
        # Determine box styling based on type
        box_style_name = f'{box_type.capitalize()} Box'
        if box_style_name not in self.document.styles:
            box_style = self.document.styles.add_style(box_style_name, WD_STYLE_TYPE.PARAGRAPH)
            box_style.font.size = Pt(11)
            box_style.paragraph_format.space_after = Pt(6)
            box_style.paragraph_format.space_before = Pt(6)
            box_style.paragraph_format.left_indent = Inches(0.25)
            box_style.paragraph_format.right_indent = Inches(0.25)

        # Create title style if needed
        title_style_name = f'{box_type.capitalize()} Box Title'
        if title_style_name not in self.document.styles:
            title_style = self.document.styles.add_style(title_style_name, WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(12)
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(6)
            title_style.paragraph_format.left_indent = Inches(0.25)

        # Add box icon based on type
        box_icon = "ℹ️ "  # Default info icon
        if box_type == "warning":
            box_icon = "⚠️ "
        elif box_type == "tip":
            box_icon = "💡 "
        elif box_type == "research":
            box_icon = "🔍 "
        elif box_type == "fact":
            box_icon = "📊 "

        # Create a table for the box with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell and set its properties
        cell = table.cell(0, 0)

        # Add title
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(f"{box_icon}{title}")
        title_run.bold = True
        title_para.style = title_style_name

        # Add content
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = box_style_name

        # Add spacing after the box
        self.document.add_paragraph()

        return table

    def add_lily_callout(self, content):
        """
        Add a simple Lily callout with the given content.
        This is used in the original research pack orchestrator.

        Args:
            content: The callout content
        """
        # Create a table for the callout with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add light blue shading to the cell background
        shading_element = cell._element.tcPr.xpath('./w:shd')
        if not shading_element:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6F0FF"/>')
            cell._element.tcPr.append(shading_element)
        else:
            shading_element[0].set(qn('w:fill'), 'E6F0FF')  # Light blue background

        # Add Lily icon and content
        content_para = cell.paragraphs[0]
        content_run = content_para.add_run(f"💡 Lily's Insight: {content}")
        content_run.italic = True
        content_run.font.size = Pt(11)

        # Add spacing after the callout
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_lily_callout(self, content):
        """
        Add a simple Lily callout with the given content.
        This is used in the original research pack orchestrator.

        Args:
            content: The callout content
        """
        # Create a table for the callout with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add light blue shading to the cell background
        shading_element = cell._element.tcPr.xpath('./w:shd')
        if not shading_element:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6F0FF"/>')
            cell._element.tcPr.append(shading_element)
        else:
            shading_element[0].set(qn('w:fill'), 'E6F0FF')  # Light blue background

        # Add Lily icon and content
        content_para = cell.paragraphs[0]
        content_run = content_para.add_run(f"💡 Lily's Insight: {content}")
        content_run.italic = True
        content_run.font.size = Pt(11)

        # Add spacing after the callout
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_lily_callout_box(self, title, content, callout_type="guidance"):
        """
        Add a special Lily callout box with visual flair

        Args:
            title: The title of the callout
            content: The main text content
            callout_type: Type of callout (guidance, research, warning, insight, tip, confidence, question, brainstorm, caution)
        """
        # Determine icon based on callout type
        icon = "✨ "  # Default icon

        # Map callout types to icons
        icon_map = {
            "research": "📚 ",
            "warning": "⚠️ ",
            "insight": "💡 ",
            "tip": "💡 ",
            "confidence": "🌟 ",
            "question": "❓ ",
            "brainstorm": "🧠 ",
            "caution": "⚠️ "
        }

        # Get the icon from the map, or use default
        if callout_type in icon_map:
            icon = icon_map[callout_type]

        # Create a table for the callout with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add title with icon
        title_para = cell.paragraphs[0]
        title_para.text = ""

        # Format the title based on callout type
        # For some callout types, we don't want to include the type in the title
        if callout_type in ["tip", "confidence", "question", "brainstorm", "insight"]:
            formatted_title = f"{icon}{title}"
        else:
            formatted_title = f"{icon}Lily's {callout_type.capitalize()}: {title}"

        title_run = title_para.add_run(formatted_title)
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add the content
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = 'Normal'

        # Add spacing after the callout
        self.document.add_paragraph()

        return table

    def add_research_highlight(self, research_title, source, content, year=None):
        """
        Add a research highlight box that showcases key research findings

        Args:
            research_title: Title of the research paper
            source: Source of the research (authors/journal)
            content: Key findings or content to highlight
            year: Publication year (optional)
        """
        # Create a table for the research highlight with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add research icon and title
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(f"🔍 Research: {research_title}")
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add source information
        source_text = source
        if year:
            source_text += f" ({year})"

        source_para = cell.add_paragraph()
        source_run = source_para.add_run(source_text)
        source_run.italic = True
        source_run.font.size = Pt(10)

        # Add the key findings
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = 'Normal'

        # Add spacing after the research highlight
        self.document.add_paragraph()

        return table

    def add_side_note(self, text, note_type="tip"):
        """
        Add a side note with different styling than the main content

        Args:
            text: The side note text
            note_type: Type of note (tip, note, important)
        """
        # Create a paragraph for the side note
        paragraph = self.document.add_paragraph()

        # Determine icon based on note type
        icon = "💡 "
        if note_type == "note":
            icon = "📝 "
        elif note_type == "important":
            icon = "❗ "

        # Add icon and text
        run = paragraph.add_run(f"{icon} {text}")
        run.italic = True

        # Style the paragraph
        if f'{note_type.capitalize()} Note' not in self.document.styles:
            note_style = self.document.styles.add_style(f'{note_type.capitalize()} Note', WD_STYLE_TYPE.PARAGRAPH)
            note_style.font.size = Pt(10)
            note_style.font.italic = True
            note_style.paragraph_format.left_indent = Inches(0.5)
            note_style.paragraph_format.right_indent = Inches(0.5)
            note_style.paragraph_format.space_before = Pt(6)
            note_style.paragraph_format.space_after = Pt(6)

        paragraph.style = f'{note_type.capitalize()} Note'

        return paragraph

    def add_visual_separator(self, separator_type="stars"):
        """
        Add a visual separator to break up content sections

        Args:
            separator_type: Type of separator (stars, line, dots)
        """
        paragraph = self.document.add_paragraph()

        # Set the separator text based on type
        if separator_type == "stars":
            separator_text = "* * * * *"
        elif separator_type == "line":
            separator_text = "―――――――――――――――――――"
        elif separator_type == "dots":
            separator_text = "• • • • •"
        else:
            separator_text = "―――――――――――――――――――"

        # Add the separator with center alignment
        run = paragraph.add_run(separator_text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing before and after
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        return paragraph

    def generate_mind_map(self, topic: str, nodes: list = None):
        """
        Generate a mind‑map image for the supplied topic using LLM to fill in the content.
        Uses llm_mind_map.py which creates a visually appealing mind map with topic-specific content.

        Returns the path to the PNG file or None if generation fails.
        """
        try:
            from llm_mind_map import generate_llm_mind_map
            return generate_llm_mind_map(topic)
        except Exception as e:
            logger.error(f"Error generating mind map with LLM: {str(e)}", exc_info=True)
            # Fall back to basic mind map if LLM version fails
            try:
                from mind_map_generator import generate_mind_map as basic_generate_mind_map
                return basic_generate_mind_map(topic, nodes)
            except Exception as e2:
                logger.error(f"Error generating fallback mind map: {str(e2)}", exc_info=True)
                return None

    def generate_research_journey_map(self, topic: str):
        """
        Generate a research journey map for the supplied topic using LLM to fill in the content.
        Uses the modular diagram_generators package which creates a visually appealing journey map with topic-specific content.

        Returns the path to the PNG file or None if generation fails.
        """
        try:
            from diagram_generators import generate_research_journey_map
            return generate_research_journey_map(topic)
        except Exception as e:
            logger.error(f"Error generating research journey map with LLM: {str(e)}", exc_info=True)
            # Fall back to basic diagram if LLM version fails
            try:
                from fix_diagrams import generate_research_process_diagram as fallback_diagram
                return fallback_diagram(topic)
            except Exception as e2:
                logger.error(f"Error generating fallback research journey map: {str(e2)}", exc_info=True)
                return None

    def generate_research_process_diagram(self, topic: str):
        """
        Generate a research process diagram for the supplied topic using LLM to fill in the content.
        Uses the modular diagram_generators package which creates a visually appealing diagram with topic-specific content.

        Returns the path to the PNG file or None if generation fails.
        """
        try:
            from diagram_generators import generate_research_process_diagram
            return generate_research_process_diagram(topic)
        except Exception as e:
            logger.error(f"Error generating research process diagram with LLM: {str(e)}", exc_info=True)
            # Fall back to basic diagram if LLM version fails
            try:
                from fix_diagrams import generate_research_process_diagram as _gen_rpd
                return _gen_rpd(topic)
            except Exception as e2:
                logger.error(f"Error generating fallback research process diagram: {str(e2)}", exc_info=True)
                return None

    def add_diagram_to_document(self, diagram_path: str, caption: str = None, use_landscape: bool = False):
        """
        Insert a diagram into the DOCX, centred and scaled to 6 inches wide.
        A numbered caption is added automatically when supplied.

        Args:
            diagram_path: Path to the diagram image file
            caption: Optional caption for the diagram
            use_landscape: Whether to use landscape orientation for the diagram

        Returns:
            True if the diagram was added successfully, False otherwise
        """
        try:
            if diagram_path and os.path.exists(diagram_path):
                # Check if this is a mind map (based on filename)
                is_mind_map = 'mind_map' in os.path.basename(diagram_path).lower()

                # Use landscape orientation for mind maps by default
                if is_mind_map:
                    use_landscape = True

                # Use landscape formatter if requested
                if use_landscape:
                    try:
                        # Import the landscape formatter
                        import sys
                        sys.path.append('/home/admin/projects/kdd')
                        from diagram_orchestrator.utils.landscape_formatter import add_landscape_diagram_page

                        # Auto‑increment figure counter
                        self.figure_count = getattr(self, "figure_count", 0) + 1

                        # Update caption with figure number
                        if caption:
                            clean_caption = caption.replace("Figure:", "").replace("##", "").strip()
                            full_caption = f"Figure {self.figure_count}. {clean_caption}"
                        else:
                            full_caption = None

                        # Add the diagram in landscape orientation
                        result = add_landscape_diagram_page(self.document, diagram_path, full_caption)
                        if result:
                            logger.info(f"Successfully added diagram in landscape orientation: {diagram_path}")
                            return True
                        else:
                            logger.warning(f"Failed to add diagram in landscape orientation, falling back to portrait: {diagram_path}")
                    except Exception as landscape_error:
                        logger.error(f"Error using landscape formatter: {str(landscape_error)}")
                        logger.info("Falling back to portrait orientation")

                # Standard portrait orientation (fallback)
                # Spacing before the figure
                self.document.add_paragraph().paragraph_format.space_before = Pt(12)

                # Add image
                para = self.document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                picture = run.add_picture(diagram_path, width=Inches(6))

                # Maintain aspect ratio if docx tracked dimensions
                if hasattr(picture, "height") and hasattr(picture, "width") and picture.width:
                    aspect_ratio = picture.height / picture.width
                    picture.width = Inches(6)
                    picture.height = int(Inches(6) * aspect_ratio)

                # Auto‑increment figure counter
                self.figure_count = getattr(self, "figure_count", 0) + 1

                # Caption handling
                if caption:
                    clean_caption = caption.replace("Figure:", "").replace("##", "").strip()
                    cap_para = self.document.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.style = "Caption"
                    cap_para.add_run(f"Figure {self.figure_count}. {clean_caption}")

                # Spacing after the figure
                self.document.add_paragraph().paragraph_format.space_after = Pt(12)

                logger.info(f"Successfully added diagram from {diagram_path}")
                return True
            else:
                logger.warning(f"Diagram file not found: {diagram_path}")
                return False
        except Exception as e:
            logger.error(f"Error adding diagram to document: {str(e)}", exc_info=True)
            return False

    def add_research_data_section(self, title, data_points, source=None):
        """
        Add a section with research data presented in a visually appealing way

        Args:
            title: Title for the research data section
            data_points: List of research findings or data points
            source: Data source (optional)
        """
        # Create a table for the data section
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add title with research icon
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(f"📊 {title}")
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add source if provided
        if source:
            source_para = cell.add_paragraph()
            source_run = source_para.add_run(f"Source: {source}")
            source_run.italic = True
            source_run.font.size = Pt(9)

        # Add each data point as bullet point
        for data_point in data_points:
            point_para = cell.add_paragraph()
            point_para.style = 'List Bullet'
            point_para.text = data_point

        # Add spacing after the data section
        self.document.add_paragraph()

        return table

    def add_argument_planning_grid(self, topic: str):
        """
        Add an argument planning grid for the given topic

        Args:
            topic: The research topic
        """
        # Create a table for the argument planning grid
        table = self.document.add_table(rows=5, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Main Argument
        table.columns[1].width = Inches(2.0)  # Supporting Evidence
        table.columns[2].width = Inches(1.5)  # Counterargument
        table.columns[3].width = Inches(1.5)  # Rebuttal

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Main Argument"
        header_cells[1].text = "Supporting Evidence"
        header_cells[2].text = "Counterargument"
        header_cells[3].text = "Rebuttal"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run(f"Use this grid to plan your arguments for your research on {topic}. Start with your strongest point and build from there.").italic = True

        # Add a Lily tip
        self.add_lily_callout_box(
            "Lily's Argument Planning Tip",
            "Start with your strongest point and build from there. Each argument deserves backup and balance. Remember to anticipate counterarguments and prepare thoughtful rebuttals.",
            "tip"
        )

        return table

    def add_emergency_research_plan(self):
        """
        Add an emergency research recovery plan
        """
        # Add heading for the emergency plan
        self.document.add_paragraph("If you're behind or overwhelmed, use this 6-step plan to get back on track fast:", style='Normal').bold = True

        # Create a table for the emergency plan
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(0.5)  # Step number
        table.columns[1].width = Inches(5.5)  # Step description

        # Add steps
        steps = [
            "Define your research question in one sentence (skip perfection).",
            "List 3 key terms and search academic sources using those only.",
            "Choose one framework or methodology — go with what's manageable.",
            "Draft a skeleton structure: Intro, Main Argument, Counterpoint.",
            "Turn your bullet points into paragraphs.",
            "Submit a solid draft — not a perfect essay."
        ]

        # Add each step to the table
        for i, step in enumerate(steps):
            row = table.rows[i].cells
            row[0].text = str(i + 1)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = step

        # Add a Lily confidence booster
        self.add_lily_callout_box(
            "Lily's Confidence Booster",
            "Progress beats panic. Always. Remember that every researcher faces time crunches. This emergency plan will help you produce something substantial even when time is tight.",
            "confidence"
        )

        return table

    def add_editable_timeline(self, topic: str):
        """
        Add an editable research timeline planner

        Args:
            topic: The research topic
        """
        # Create a table for the timeline
        table = self.document.add_table(rows=9, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(0.5)  # Week
        table.columns[1].width = Inches(1.0)  # Goal
        table.columns[2].width = Inches(3.5)  # Tasks
        table.columns[3].width = Inches(1.0)  # Done

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Week"
        header_cells[1].text = "Goal"
        header_cells[2].text = "Tasks"
        header_cells[3].text = "Done (✓)"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add timeline content
        timeline_content = [
            ["1", "Topic Selection", "Brainstorm, choose topic, refine question", ""],
            ["2", "Literature Review", "Collect 5-10 sources, start reading", ""],
            ["3", "Methodology Choice", "Decide on approach, start outline", ""],
            ["4", "Drafting", "Intro + 1st Argument", ""],
            ["5", "Drafting", "Continue body + analysis", ""],
            ["6", "Drafting", "Write conclusion + integrate citations", ""],
            ["7", "Editing", "Proofreading, peer review", ""],
            ["8", "Finalise", "Add references, polish formatting, submit", ""]
        ]

        # Add each row to the table
        for i, row_content in enumerate(timeline_content):
            row = table.rows[i+1].cells
            for j, cell_content in enumerate(row_content):
                row[j].text = cell_content
                if j == 0:  # Center the week number
                    row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run(f"Use this timeline to plan your research on {topic}. Adjust the weeks and tasks to fit your schedule and deadlines.").italic = True

        return table

    def add_reflection_prompts(self, topic: str):
        """
        Add reflection and self-assessment prompts

        Args:
            topic: The research topic
        """
        # Create a table for the reflection log
        table = self.document.add_table(rows=5, cols=5)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(0.8)  # Date
        table.columns[1].width = Inches(1.5)  # What I Worked On
        table.columns[2].width = Inches(1.2)  # What Worked
        table.columns[3].width = Inches(1.2)  # What Was Hard
        table.columns[4].width = Inches(1.3)  # What I'll Try Next

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Date"
        header_cells[1].text = "What I Worked On"
        header_cells[2].text = "What Worked"
        header_cells[3].text = "What Was Hard"
        header_cells[4].text = "What I'll Try Next"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run(f"Use this reflection log to track your progress and learning as you research {topic}. Regular reflection helps identify what's working and what needs adjustment.").italic = True

        # Add reflection prompts
        self.add_lily_callout_box(
            "Lily's Reflection Prompts",
            "Consider these questions as you reflect on your research process:\n\n" +
            "• What new insights did I gain today?\n" +
            "• Which sources were most helpful and why?\n" +
            "• What connections am I seeing between different aspects of my topic?\n" +
            "• Where am I feeling stuck or confused?\n" +
            "• What's one thing I could do differently tomorrow?",
            "question"
        )

        return table

    def add_glossary_builder(self, topic: str):
        """
        Add a glossary builder template

        Args:
            topic: The research topic
        """
        # Create a table for the concept clarifier
        table = self.document.add_table(rows=6, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Concept
        table.columns[1].width = Inches(2.0)  # My Explanation
        table.columns[2].width = Inches(2.0)  # Academic Definition
        table.columns[3].width = Inches(0.5)  # Source

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Concept"
        header_cells[1].text = "My Explanation"
        header_cells[2].text = "Academic Definition"
        header_cells[3].text = "Source"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run(f"Use this glossary builder to clarify key concepts related to {topic}. Writing concepts in your own words helps solidify your understanding.").italic = True

        # Add a Lily tip
        self.add_lily_callout_box(
            "Lily's Learning Tip",
            "If you can define it in your own words, you understand it. If you can teach it, you *own* it. Try explaining these concepts to someone else to test your understanding.",
            "tip"
        )

        return table

    def add_peer_review_checklist(self):
        """
        Add a peer review checklist
        """
        # Create a table for the checklist
        table = self.document.add_table(rows=10, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(5.0)  # Checklist item
        table.columns[1].width = Inches(1.0)  # Yes/No

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Review Criteria"
        header_cells[1].text = "Yes/No"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add checklist items
        checklist_items = [
            "Does the introduction clearly state the research question or thesis?",
            "Is the literature review comprehensive and well-organized?",
            "Is the methodology appropriate for the research question?",
            "Are arguments supported with evidence and examples?",
            "Are counterarguments addressed effectively?",
            "Is the conclusion strong and connected to the introduction?",
            "Are all sources properly cited in the required format?",
            "Is the writing clear, concise, and free of errors?",
            "Does the paper flow logically from one section to the next?"
        ]

        # Add each item to the table
        for i, item in enumerate(checklist_items):
            row = table.rows[i+1].cells
            row[0].text = item

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run("Use this checklist when reviewing your own work or when asking a peer to review your paper. Constructive feedback is essential for improving your research.").italic = True

        return table

    def add_source_tracking_table(self):
        """
        Add a source tracking table
        """
        # Create a table for tracking sources
        table = self.document.add_table(rows=6, cols=5)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Source Title
        table.columns[1].width = Inches(1.0)  # Author
        table.columns[2].width = Inches(0.5)  # Year
        table.columns[3].width = Inches(2.0)  # Key Argument
        table.columns[4].width = Inches(0.5)  # Usefulness

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Source Title"
        header_cells[1].text = "Author"
        header_cells[2].text = "Year"
        header_cells[3].text = "Key Argument"
        header_cells[4].text = "Usefulness (1-5)"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a Lily tip
        self.add_lily_callout_box(
            "Lily's Source Evaluation Tip",
            "Strong sources make strong writing. Evaluate as you collect. Consider factors like author credentials, publication date, peer review status, and relevance to your specific topic.",
            "tip"
        )

        return table

    def add_introduction_builder(self):
        """
        Add an introduction builder template
        """
        # Create a table for the introduction builder
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Element
        table.columns[1].width = Inches(4.0)  # Notes

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Element"
        header_cells[1].text = "Notes"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add introduction elements
        elements = [
            "Hook / Opener",
            "Topic Definition",
            "Scope & Boundaries",
            "Why It Matters",
            "Thesis Statement"
        ]

        # Add each element to the table
        for i, element in enumerate(elements):
            row = table.rows[i+1].cells
            row[0].text = element
            row[0].paragraphs[0].runs[0].bold = True

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run("Use this template to plan your introduction. A strong introduction sets the stage for your entire paper and engages your reader from the start.").italic = True

        # Add a Lily tip
        self.add_lily_callout_box(
            "Lily's Introduction Tip",
            "Think of this as your elevator pitch to an academic. Your introduction should clearly communicate what your paper is about, why it matters, and what you'll argue. Write it last, even though it appears first!",
            "tip"
        )

        return table

    def add_paragraph_planner(self):
        """
        Add a paragraph planner template using the PEEL format
        """
        # Create a table for the paragraph planner
        table = self.document.add_table(rows=4, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Point
        table.columns[1].width = Inches(1.5)  # Evidence
        table.columns[2].width = Inches(1.5)  # Explanation
        table.columns[3].width = Inches(1.5)  # Link

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Point"
        header_cells[1].text = "Evidence"
        header_cells[2].text = "Explanation"
        header_cells[3].text = "Link"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run("Use this PEEL (Point, Evidence, Explanation, Link) format to plan your paragraphs. This structure helps ensure each paragraph has a clear purpose and supports your overall argument.").italic = True

        # Add a Lily tip explaining PEEL
        self.add_lily_callout_box(
            "Lily's PEEL Format Guide",
            "P - Point: Start with your main idea or claim\n" +
            "E - Evidence: Support your point with evidence (quotes, data, examples)\n" +
            "E - Explanation: Explain how your evidence supports your point\n" +
            "L - Link: Connect back to your thesis or transition to the next paragraph",
            "tip"
        )

        return table

    def add_thesis_alignment_check(self):
        """
        Add a thesis alignment check template
        """
        # Create a table for the thesis alignment check
        table = self.document.add_table(rows=5, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Central Claim
        table.columns[1].width = Inches(1.0)  # Supporting?
        table.columns[2].width = Inches(2.0)  # How?
        table.columns[3].width = Inches(1.0)  # Keep It?

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Central Claim"
        header_cells[1].text = "Supporting? (Y/N)"
        header_cells[2].text = "How?"
        header_cells[3].text = "Keep It?"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run("Use this table to check if all your arguments align with your thesis. Every claim in your paper should support your central argument in some way.").italic = True

        # Add a Lily insight
        self.add_lily_callout_box(
            "Lily's Thesis Alignment Insight",
            "A common issue in academic writing is including interesting information that doesn't actually support your thesis. This table helps you identify and remove tangents that might distract from your main argument.",
            "insight"
        )

        return table

    def add_citation_cheat_sheet(self, citation_style="Harvard"):
        """
        Add a citation cheat sheet for the specified style

        Args:
            citation_style: The citation style (Harvard, APA, MLA, Chicago)
        """
        # Create a table for the citation cheat sheet
        table = self.document.add_table(rows=5, cols=1)
        table.style = 'Light Grid'

        # Get the cell for the title
        title_cell = table.cell(0, 0)
        title_para = title_cell.paragraphs[0]
        title_run = title_para.add_run(f"{citation_style} Citation Cheat Sheet")
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add citation examples based on style
        if citation_style == "Harvard":
            # In-text citation
            intext_cell = table.cell(1, 0)
            intext_para = intext_cell.paragraphs[0]
            intext_para.add_run("In-text Citation Example:\n").bold = True
            intext_para.add_run("(Smith, 2020) or Smith (2020) argues that...")

            # Book reference
            book_cell = table.cell(2, 0)
            book_para = book_cell.paragraphs[0]
            book_para.add_run("Book Reference Format:\n").bold = True
            book_para.add_run("Smith, J. (2020). Book Title. Publisher.")

            # Journal article reference
            journal_cell = table.cell(3, 0)
            journal_para = journal_cell.paragraphs[0]
            journal_para.add_run("Journal Article Format:\n").bold = True
            journal_para.add_run("Smith, J. (2020). Article title. Journal Name, 12(3), 45-67.")

            # Website reference
            website_cell = table.cell(4, 0)
            website_para = website_cell.paragraphs[0]
            website_para.add_run("Website Format:\n").bold = True
            website_para.add_run("Johnson, A. (2023). Article Title. [online] Available at: URL [Accessed 18 Apr. 2025].")

        elif citation_style == "APA":
            # Similar structure for APA style
            # In-text citation
            intext_cell = table.cell(1, 0)
            intext_para = intext_cell.paragraphs[0]
            intext_para.add_run("In-text Citation Example:\n").bold = True
            intext_para.add_run("(Smith, 2020) or Smith (2020) argues that...")

            # Book reference
            book_cell = table.cell(2, 0)
            book_para = book_cell.paragraphs[0]
            book_para.add_run("Book Reference Format:\n").bold = True
            book_para.add_run("Smith, J. (2020). Book title. Publisher.")

            # Journal article reference
            journal_cell = table.cell(3, 0)
            journal_para = journal_cell.paragraphs[0]
            journal_para.add_run("Journal Article Format:\n").bold = True
            journal_para.add_run("Smith, J. (2020). Article title. Journal Name, 12(3), 45-67. https://doi.org/10.xxxx")

            # Website reference
            website_cell = table.cell(4, 0)
            website_para = website_cell.paragraphs[0]
            website_para.add_run("Website Format:\n").bold = True
            website_para.add_run("Johnson, A. (2023, April 18). Article title. Website Name. https://url.com")

        # Add a Lily tip
        self.add_lily_callout_box(
            "Lily's Citation Tip",
            "Consistency is key with citations. Choose one style and stick with it throughout your paper. When in doubt, check the official style guide or use a citation management tool like Zotero or Mendeley.",
            "tip"
        )

        return table

    def add_one_week_crash_plan(self):
        """
        Add a one-week crash plan for emergency research situations
        """
        # Create a table for the crash plan
        table = self.document.add_table(rows=8, cols=4)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(0.5)  # Day
        table.columns[1].width = Inches(1.5)  # Task
        table.columns[2].width = Inches(2.0)  # Focus
        table.columns[3].width = Inches(1.0)  # Time

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Day"
        header_cells[1].text = "Task"
        header_cells[2].text = "Focus"
        header_cells[3].text = "Time"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add crash plan content
        crash_plan = [
            ["Mon", "Review Topic + Framework", "Planning", "2h"],
            ["Tue", "Research 3 Sources", "Lit Review", "3h"],
            ["Wed", "Structure Draft", "Organisation", "2h"],
            ["Thu", "Write Main Argument", "Drafting", "3h"],
            ["Fri", "Write Counterargument", "Drafting", "2h"],
            ["Sat", "Review & Add References", "Editing", "2h"],
            ["Sun", "Final Proof & Submission", "Wrap-up", "1h"]
        ]

        # Add each row to the table
        for i, row_content in enumerate(crash_plan):
            row = table.rows[i+1].cells
            for j, cell_content in enumerate(row_content):
                row[j].text = cell_content
                if j == 0:  # Center the day
                    row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a description paragraph
        self.document.add_paragraph()
        description = self.document.add_paragraph()
        description.add_run("Use this one-week crash plan when you're in panic mode and need to complete a research paper quickly. Focus on one day at a time.").italic = True

        # Add a Lily confidence booster
        self.add_lily_callout_box(
            "Lily's Reminder",
            "One day at a time. You're closer than you think. This plan is designed to help you produce a solid paper in just one week. It won't be perfect, but it will be complete and coherent.",
            "confidence"
        )

        return table
