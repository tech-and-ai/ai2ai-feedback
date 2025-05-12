import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from .table_generator import TableGenerator
import re
from app.membership.config import get_tier_config
from app.services.paper_generator.config import get_section_word_count
from typing import Dict
import textwrap
import io
import string
import random
import unicodedata

logger = logging.getLogger(__name__)

class DocumentFormatter:
    def __init__(self):
        self.document = Document()
        self._setup_document_styles(self.document)
        self._setup_sections()
        self.table_generator = TableGenerator()
        # Add section counters for proper numbering
        self.section_counters = {1: 0, 2: 0, 3: 0}

    def _setup_document_styles(self, doc: Document):
        """Setup document styles for a professional academic paper"""
        # Normal text style
        if 'Normal' in doc.styles:
            normal_style = doc.styles['Normal']
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.space_after = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5

        # Use built-in heading styles
        if 'Heading 1' in doc.styles:
            heading1 = doc.styles['Heading 1']
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.paragraph_format.space_before = Pt(24)
            heading1.paragraph_format.space_after = Pt(12)

        if 'Heading 2' in doc.styles:
            heading2 = doc.styles['Heading 2']
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.paragraph_format.space_before = Pt(18)
            heading2.paragraph_format.space_after = Pt(6)

        if 'Heading 3' in doc.styles:
            heading3 = doc.styles['Heading 3']
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.paragraph_format.space_before = Pt(12)
            heading3.paragraph_format.space_after = Pt(6)

        # Title style
        if 'Title' in doc.styles:
            title_style = doc.styles['Title']
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)

        # TOC styles
        if 'TOC Heading' not in doc.styles:
            toc_style = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(16)
            toc_style.font.bold = True
            toc_style.paragraph_format.space_before = Pt(24)
            toc_style.paragraph_format.space_after = Pt(12)

        if 'TOC' not in doc.styles:
            toc_style = doc.styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(12)
            toc_style.paragraph_format.left_indent = Inches(0.25)
            toc_style.paragraph_format.space_after = Pt(6)

        if 'TOC 1' not in doc.styles:
            toc1_style = doc.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
            toc1_style.font.size = Pt(12)
            toc1_style.paragraph_format.left_indent = Inches(0.25)
            toc1_style.paragraph_format.space_after = Pt(6)

        if 'TOC 2' not in doc.styles:
            toc2_style = doc.styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
            toc2_style.font.size = Pt(11)
            toc2_style.paragraph_format.left_indent = Inches(0.5)
            toc2_style.paragraph_format.space_after = Pt(6)

    def _setup_sections(self):
        """Setup sections for headers and footers"""
        # Create a table of contents style
        if 'TOC Heading' not in self.document.styles:
            toc_style = self.document.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(16)
            toc_style.font.bold = True
            toc_style.paragraph_format.space_before = Pt(24)
            toc_style.paragraph_format.space_after = Pt(12)

        if 'TOC 1' not in self.document.styles:
            toc1_style = self.document.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
            toc1_style.font.size = Pt(12)
            toc1_style.paragraph_format.left_indent = Inches(0.25)
            toc1_style.paragraph_format.space_after = Pt(6)

        if 'TOC 2' not in self.document.styles:
            toc2_style = self.document.styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
            toc2_style.font.size = Pt(11)
            toc2_style.paragraph_format.left_indent = Inches(0.5)
            toc2_style.paragraph_format.space_after = Pt(6)

    def create_section_divider(self, section_number: int, section_title: str, section_description: str, subsections: list, education_level: str = "university"):
        """Create a section divider page."""
        # Section colors
        section_colors = {
            1: "3498db",  # Blue
            2: "2ecc71",  # Green
            3: "e67e22",  # Orange
            4: "9b59b6",  # Purple
            5: "f1c40f",  # Yellow
        }

        # Add page break
        self.document.add_page_break()

        # Add section number and title
        section_heading = self.document.add_paragraph()
        section_heading.style = 'Heading 1'
        section_run = section_heading.add_run(f"SECTION {section_number}: {section_title}")
        section_run.font.color.rgb = RGBColor.from_string(section_colors.get(section_number, "000000"))
        section_run.font.size = Pt(24)
        section_run.font.bold = True

        # Add horizontal line
        self.document.add_paragraph("_______________________________________________")

        # Add section description
        description_para = self.document.add_paragraph()
        description_para.add_run(section_description).italic = True

        # Add subsections list
        self.document.add_paragraph("In this section:").bold = True
        for subsection in subsections:
            subsection_para = self.document.add_paragraph()
            subsection_para.style = 'List Bullet'
            subsection_para.add_run(subsection).bold = True

        # Add horizontal line
        self.document.add_paragraph("_______________________________________________")

        # Add section footer
        footer_para = self.document.add_paragraph()
        footer_para.add_run(f"Section {section_number} of 5 | {education_level.title()} Research Pack").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create_title_page(self, title: str, author: str, institution: str, date: str, education_level: str = None, logo_path: str = None):
        """Create a professional title page with Lily logo"""
        # Add Lily logo at the top
        try:
            # Use provided logo_path if available, otherwise use default path
            if not logo_path:
                logo_path = "app/static/lilylogo.png"

            if os.path.exists(logo_path):
                logo_para = self.document.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2))
                logger.info(f"Added Lily logo from {logo_path}")
            else:
                logger.warning(f"Lily logo not found at {logo_path}")
        except Exception as e:
            logger.error(f"Error adding Lily logo: {str(e)}")

        # Title - use the actual title, not the topic
        title_para = self.document.add_paragraph(title)
        title_para.style = 'Title'

        # Add some spacing
        self.document.add_paragraph()

        # Make sure we have a valid author name (not empty or None)
        if not author or author == "Researcher" or author.strip() == "":
            author = "Researcher"
            logger.warning("No valid author name provided, using 'Researcher' as default")

        # Author information - use the actual name, not "Researcher"
        author_para = self.document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"Prepared for\n{author}").bold = True

        # Institution
        inst_para = self.document.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_para.add_run(institution)

        # Date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(date)

        # Education Level (if provided)
        if education_level:
            # Format the education level for display (e.g., "high-school" -> "High School Level")
            formatted_level = education_level.replace('-', ' ').title() + " Level"

            # Add education level with distinctive formatting
            level_para = self.document.add_paragraph()
            level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            level_run = level_para.add_run(formatted_level)
            level_run.bold = True
            level_run.font.size = Pt(14)

            # Add a colored box around the education level
            level_para.paragraph_format.border_between = True
            level_para.paragraph_format.border_width = Pt(1)

            # Add a note about the education level
            note_para = self.document.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_para.add_run("This document has been adapted to the specified education level.").italic = True

        # Add page break after title page
        self.document.add_page_break()

    def add_table_of_contents(self):
        """Add a static table of contents"""
        # Add a heading for the TOC
        toc_heading = self.document.add_heading("Table of Contents", level=1)
        toc_heading.style = 'Heading 1'
        # Don't number the TOC
        toc_heading._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 0

        # Create a static table of contents with common sections
        toc_sections = [
            ("1. Introduction", 1),
            ("1.1 About This Research Pack", 2),
            ("1.2 How to Use This Document", 2),
            ("1.3 What's Included in This Pack", 2),
            ("1.4 Research Journey Map", 2),
            ("2. Research Planning Tools", 1),
            ("2.1 Four-Week Research Plan", 2),
            ("2.2 Topic Mind Map", 2),
            ("2.3 Thesis Builder", 2),
            ("2.4 Literature Review Matrix", 2),
            ("3. Topic Analysis", 1),
            ("3.1 Key Concepts", 2),
            ("3.2 Theoretical Frameworks", 2),
            ("3.3 Current Debates", 2),
            ("4. Methodological Approaches", 1),
            ("4.1 Research Methods", 2),
            ("4.2 Data Collection", 2),
            ("4.3 Analysis Techniques", 2),
            ("5. Personalized Questions", 1),
            ("6. Key Subject Matter Experts", 1),
            ("7. Citation Guide and Key Resources", 1),
            ("8. Appendices: Bonus Templates Collection", 1)
        ]

        # Add each TOC entry with appropriate style and indentation
        for section, level in toc_sections:
            toc_entry = self.document.add_paragraph()

            # Set style based on level
            if level == 1:
                toc_entry.style = 'TOC 1'
            else:
                toc_entry.style = 'TOC 2'
                toc_entry.paragraph_format.left_indent = Inches(0.25)

            # Add the section text
            toc_entry.add_run(section)

            # Add tab and page number placeholder
            toc_entry.add_run("\t")
            toc_entry.add_run("...").bold = False

        # Add a page break after TOC
        self.document.add_page_break()

    def clean_markdown(self, text: str) -> str:
        """
        Clean markdown syntax from text while preserving Lily callout tags
        """
        if not text:
            return ""

        # Check if the text contains Lily callout tags
        has_lily_callouts = "[[LILY_CALLOUT" in text

        if has_lily_callouts:
            # Temporarily replace Lily callout tags with placeholders to preserve them
            # This ensures we don't lose the square brackets around the tags
            lily_callout_pattern = r'\[\[LILY_CALLOUT.*?\]\](.*?)\[\[/LILY_CALLOUT\]\]'
            lily_callouts = re.findall(lily_callout_pattern, text, re.DOTALL)

            # Replace each callout with a unique placeholder
            placeholders = {}
            for i, match in enumerate(re.finditer(lily_callout_pattern, text, re.DOTALL)):
                placeholder = f"__LILY_CALLOUT_PLACEHOLDER_{i}__"
                placeholders[placeholder] = match.group(0)
                text = text.replace(match.group(0), placeholder, 1)

            logger.info(f"FORMATTER: Preserved {len(placeholders)} Lily callouts during markdown cleaning")

        # Remove heading markers more aggressively
        text = text.strip()
        while '#' in text:
            text = text.replace('#', '')

        # Clean other markdown syntax
        replacements = {
            '**': '',
            '__': '',
            '*': '',
            '_': '',
            '```': '',
            '`': '',
            'Figure:': '',
            'figure:': '',
            '  ': ' '  # Replace double spaces with single
        }

        # Only remove square brackets if we're not dealing with Lily callouts
        if not has_lily_callouts:
            replacements['['] = ''
            replacements[']'] = ''

        for old, new in replacements.items():
            text = text.replace(old, new)

        # Clean up multiple spaces and lines
        text = ' '.join(text.split())

        # Restore Lily callout tags if we preserved any
        if has_lily_callouts:
            for placeholder, original in placeholders.items():
                text = text.replace(placeholder, original)

            logger.info(f"FORMATTER: Restored {len(placeholders)} Lily callouts after markdown cleaning")

        return text.strip()

    def add_section(self, heading: str, content: str, level: int = 1):
        """Add a new section with heading and content"""
        # Add page break before each main section (level 1)
        if level == 1 and self.document.paragraphs:
            self.document.add_page_break()

        # Clean up markdown-style headings
        heading = self.clean_markdown(heading)
        heading = heading.lstrip('#').strip()

        # Handle section numbering
        if level <= 3:  # We only number up to level 3
            if level == 1:
                # Increment level 1 counter, reset others
                self.section_counters[1] += 1
                self.section_counters[2] = 0
                self.section_counters[3] = 0
                section_number = f"{self.section_counters[1]}."
            elif level == 2:
                # Increment level 2 counter, reset level 3
                self.section_counters[2] += 1
                self.section_counters[3] = 0
                section_number = f"{self.section_counters[1]}.{self.section_counters[2]}"
            elif level == 3:
                # Increment level 3 counter
                self.section_counters[3] += 1
                section_number = f"{self.section_counters[1]}.{self.section_counters[2]}.{self.section_counters[3]}"

            # Add numbered heading (except for special sections)
            special_sections = ["references", "abstract", "key points", "appendix", "acknowledgements", "table of contents"]
            if heading.lower() not in special_sections:
                heading_with_number = f"{section_number} {heading}"
            else:
                heading_with_number = heading
                if level == 1:
                    self.section_counters[1] -= 1
        else:
            heading_with_number = heading

        # Add extra spacing before section heading
        if level == 1:
            self.document.add_paragraph().paragraph_format.space_before = Pt(24)

        # Add the section heading with appropriate style
        section_heading = self.document.add_heading(heading_with_number, level=level)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply the proper heading style
        if level == 1:
            section_heading.style = 'Heading 1'
        elif level == 2:
            section_heading.style = 'Heading 2'
        elif level == 3:
            section_heading.style = 'Heading 3'
        else:
            section_heading.style = 'Heading 1'

        # Process content and add paragraphs with proper formatting
        if content:
            # Check for visual journey map tag
            visual_journey_map_match = re.search(r'\[VISUAL_JOURNEY_MAP\](.*?)\[\/VISUAL_JOURNEY_MAP\]', content)
            if visual_journey_map_match:
                # Extract the image path
                image_path = visual_journey_map_match.group(1).strip()
                # Remove the tag from the content
                content = content.replace(visual_journey_map_match.group(0), '')
                # Add the image if it exists
                if os.path.exists(image_path):
                    try:
                        # Add the journey map image
                        self.document.add_paragraph().paragraph_format.space_before = Pt(12)
                        paragraph = self.document.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(6))
                        logger.info(f"Added research journey map image from {image_path}")
                        # Add caption
                        caption_para = self.document.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_para.add_run("Figure: Research Journey Map").italic = True
                        self.document.add_paragraph().paragraph_format.space_after = Pt(12)
                    except Exception as e:
                        logger.error(f"Error adding journey map image: {str(e)}")

            # Split content into paragraphs
            paragraphs = content.split('\n\n')

            # Add each paragraph with proper spacing
            for i, para_text in enumerate(paragraphs):
                if para_text.strip():
                    # Clean the paragraph text
                    para_text = self.clean_markdown(para_text)

                    # Add the paragraph
                    paragraph = self.document.add_paragraph()
                    paragraph.text = para_text.strip()

                    # Set paragraph formatting
                    paragraph.style = 'Normal'
                    paragraph.paragraph_format.first_line_indent = Inches(0.5)
                    paragraph.paragraph_format.space_after = Pt(12)

                    # Add extra spacing between subsections if detected
                    if i > 0 and para_text.strip().startswith(str(self.section_counters[1])):
                        paragraph.paragraph_format.space_before = Pt(18)

        # Add extra spacing after the section
        if level == 1:
            self.document.add_paragraph().paragraph_format.space_after = Pt(24)

    def add_bullet_point(self, text: str):
        """Add a bullet point to the document with proper formatting"""
        paragraph = self.document.add_paragraph(style='List Bullet')
        paragraph.text = text.strip()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        return paragraph

    def add_heading(self, text: str, level: int = 1):
        """Add a heading with the specified level"""
        if not text:
            return None

        # Clean markdown formatting
        text = self.clean_markdown(text)

        # Create the heading
        heading = self.document.add_heading(text, level=level)

        # Apply appropriate formatting based on level
        if level == 0:  # Major section divider
            heading.style = self.document.styles['Title']
            heading.paragraph_format.space_before = Pt(24)
            heading.paragraph_format.space_after = Pt(12)
        elif level == 1:
            heading.style = self.document.styles['Heading 1']
        elif level == 2:
            heading.style = self.document.styles['Heading 2']
        elif level == 3:
            heading.style = self.document.styles['Heading 3']

        return heading

    def add_paragraph(self, text: str = ""):
        """Add a paragraph with the specified text, processing any embedded Lily insights"""
        if not text:
            return self.document.add_paragraph()

        # Check for Lily callout markers before cleaning markdown
        if "[[LILY_CALLOUT" in text:
            logger.info(f"FORMATTER: Found [[LILY_CALLOUT marker in text before cleaning")

            # Use a direct approach to process Lily callouts
            # This is a simplified version that just looks for the markers and processes them

            # Define patterns for both formats
            pattern_original = r'\[\[LILY_CALLOUT\s+type="(\w+)"\s+title="([^"]+)"(?:\s+source="([^"]+)")?\s*\]\]([\s\S]*?)\[\[\/LILY_CALLOUT\]\]'
            pattern_alt = r'\[\[LILY_CALLOUT:([^:]+):([^\]]+)\]\](.*?)\[\[/LILY_CALLOUT\]\]'

            # Find all matches for both patterns
            matches_original = list(re.finditer(pattern_original, text, re.DOTALL))
            matches_alt = list(re.finditer(pattern_alt, text, re.DOTALL))

            # Combine all matches and sort by position
            all_matches = []
            for match in matches_original:
                all_matches.append({
                    'start': match.start(),
                    'end': match.end(),
                    'type': match.group(1),
                    'title': match.group(2),
                    'source': match.group(3) or "",
                    'content': match.group(4)
                })

            for match in matches_alt:
                all_matches.append({
                    'start': match.start(),
                    'end': match.end(),
                    'type': match.group(2),
                    'title': match.group(1),
                    'source': "",
                    'content': match.group(3)
                })

            # Sort matches by position
            all_matches.sort(key=lambda x: x['start'])

            logger.info(f"FORMATTER: Found {len(all_matches)} Lily callouts in text")

            # Process the text with callouts
            if all_matches:
                # Add text before first callout
                if all_matches[0]['start'] > 0:
                    before_text = text[:all_matches[0]['start']].strip()
                    if before_text:
                        before_para = self.document.add_paragraph()
                        before_para.text = self.clean_markdown(before_text)
                        before_para.style = 'Normal'

                # Process each callout and text between callouts
                for i, match in enumerate(all_matches):
                    # Add the callout
                    self.add_lily_callout_box(match['title'], match['content'], match['type'], source=match['source'])

                    # Add text after this callout and before the next one (or end of text)
                    if i < len(all_matches) - 1:
                        between_text = text[match['end']:all_matches[i+1]['start']].strip()
                    else:
                        between_text = text[match['end']:].strip()

                    if between_text:
                        between_para = self.document.add_paragraph()
                        between_para.text = self.clean_markdown(between_text)
                        between_para.style = 'Normal'

                # Return the last paragraph added
                return self.document.paragraphs[-1]

        # If no Lily callout markers found or if we want to use the original approach,
        # clean markdown formatting and proceed with the normal flow
        text = self.clean_markdown(text)

        # Log the text being processed for Lily callouts
        logger.info(f"FORMATTER: Processing text for Lily callouts (length: {len(text)} chars)")

        # Check if the text contains any Lily callout markers after cleaning
        if "[[LILY_CALLOUT" in text:
            logger.info(f"FORMATTER: Found [[LILY_CALLOUT marker in text after cleaning: {text[:200]}...")
            # Print the exact text for debugging
            logger.info(f"FORMATTER: Full text: {repr(text)}")
        else:
            logger.info("FORMATTER: NO [[LILY_CALLOUT marker found in text")

        # Add as normal paragraph if no special formatting needed
        paragraph = self.document.add_paragraph()
        paragraph.text = text
        paragraph.style = 'Normal'
        return paragraph

    def add_page_break(self):
        """Add a page break to the document"""
        self.document.add_page_break()

    def add_figure(self, image_path: str, caption: str):
        """Add a figure with caption and maintain aspect ratio"""
        try:
            logger.info(f"FIGURE DEBUG: add_figure called with image_path={image_path}, caption={caption}")

            if not image_path:
                logger.error(f"FIGURE DEBUG: image_path is None or empty")
                return

            if not os.path.exists(image_path):
                logger.error(f"FIGURE DEBUG: Image not found at path: {image_path}")
                # Try to find the file in a different location
                base_filename = os.path.basename(image_path)
                logger.info(f"FIGURE DEBUG: Trying to find {base_filename} in alternative locations")

                # Check in the generated_diagrams directory
                alt_path = f"/home/admin/projects/kdd/essential-apis/clean-api/app/static/generated_diagrams/{base_filename}"
                if os.path.exists(alt_path):
                    logger.info(f"FIGURE DEBUG: Found image at alternative path: {alt_path}")
                    image_path = alt_path
                else:
                    logger.error(f"FIGURE DEBUG: Could not find image in alternative locations")
                    return

            # Add spacing before figure
            self.document.add_paragraph().paragraph_format.space_before = Pt(12)
            logger.info(f"FIGURE DEBUG: Added spacing before figure")

            # Center the figure
            figure_para = self.document.add_paragraph()
            figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure_run = figure_para.add_run()

            try:
                logger.info(f"FIGURE DEBUG: Attempting to add picture from {image_path}")
                picture = figure_run.add_picture(image_path)
                logger.info(f"FIGURE DEBUG: Successfully added picture to document")

                # Set width while maintaining aspect ratio
                target_width = Inches(6.0)  # Standard width for figures
                aspect_ratio = picture.height / picture.width
                picture.width = target_width
                picture.height = int(target_width * aspect_ratio)
                logger.info(f"FIGURE DEBUG: Set picture dimensions: width={target_width}, height={int(target_width * aspect_ratio)}")
            except Exception as pic_error:
                logger.error(f"FIGURE DEBUG: Error adding picture: {str(pic_error)}")
                return

            # Add figure number
            self.figure_count = getattr(self, 'figure_count', 0) + 1
            logger.info(f"FIGURE DEBUG: Incremented figure count to {self.figure_count}")

            # Create or get the Caption style
            if 'Caption' not in self.document.styles:
                caption_style = self.document.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                caption_style.font.size = Pt(10)
                caption_style.font.italic = False
                caption_style.font.bold = True
                caption_style.paragraph_format.space_before = Pt(6)
                caption_style.paragraph_format.space_after = Pt(18)
                caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.info(f"FIGURE DEBUG: Created Caption style")

            # Clean up caption and add with automatic numbering
            clean_caption = caption.replace('Figure:', '').strip()  # Remove any existing 'Figure:' prefix
            clean_caption = clean_caption.replace('##', '').strip()  # Remove markdown heading syntax
            caption_para = self.document.add_paragraph()
            caption_para.style = 'Caption'
            caption_para.add_run(f'Figure {self.figure_count}. {self.clean_markdown(caption)}')
            logger.info(f"FIGURE DEBUG: Added caption: 'Figure {self.figure_count}. {self.clean_markdown(caption)}'")

            # Add spacing after figure
            self.document.add_paragraph().paragraph_format.space_after = Pt(12)
            logger.info(f"FIGURE DEBUG: Added spacing after figure")

            logger.info(f"FIGURE DEBUG: Successfully added figure from {image_path}")
            return True
        except Exception as e:
            logger.error(f"FIGURE DEBUG: Error adding figure: {str(e)}")
            logger.error(f"FIGURE DEBUG: Image path: {image_path}")
            return False

    def add_header_footer(self, title: str):
        """Add header and footer with page numbers and copyright text to the document"""
        # Default footer text - can be modified here for easy updates
        FOOTER_TEXT = "© 2025 researchassistant.uk - All rights reserved"

        # Get the section
        section = self.document.sections[0]

        # Add title to header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = title
        header_para.style = self.document.styles['Normal']
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page numbers and copyright to footer
        footer = section.footer
        footer_para = footer.paragraphs[0]

        # Add random Lily encouragement statement
        encouragements = [
            # Original 15 encouragements
            "Lily is here to help you succeed in your academic journey!",
            "Remember, every great researcher started with a single question. Keep going!",
            "You're making real progress—don't be afraid to explore new ideas!",
            "Every challenge is a chance to learn. Stay curious!",
            "You've got this! Your unique perspective matters.",
            "Keep pushing forward—breakthroughs come from persistence!",
            "Celebrate your progress—each step brings you closer to your goal!",
            "Take a moment to appreciate how far you've come already!",
            "Research is a journey, not a destination. Enjoy the process!",
            "Your ideas matter. Keep developing them!",
            "Small steps lead to big discoveries. You're on the right track!",
            "Struggling? That's part of the process. Keep going!",
            "Your curiosity is your superpower. Use it wisely!",
            "Great research starts with great questions. You're asking the right ones!",
            "Remember, clarity beats complexity. You're doing great!",

            # Additional 30 encouragements
            "The best research comes from genuine curiosity. Follow yours!",
            "Academic excellence is built one paragraph at a time.",
            "Your perspective is valuable—don't be afraid to share it!",
            "Research is like a puzzle—each piece matters to the whole picture.",
            "Today's notes become tomorrow's breakthroughs.",
            "Learning to ask good questions is half the battle in research.",
            "When you feel stuck, take a step back and look at the bigger picture.",
            "The most interesting discoveries often happen when you least expect them.",
            "Your academic voice matters—keep developing it!",
            "Research skills improve with practice—you're getting better every day.",
            "Critical thinking is a muscle—you're strengthening it right now.",
            "Good research takes time—be patient with yourself.",
            "Each draft brings you closer to clarity. Keep refining!",
            "The best researchers are also the most persistent ones.",
            "Your unique approach to this topic is what makes your work valuable.",
            "Academic growth happens outside your comfort zone—you're doing great!",
            "Remember to celebrate small victories in your research journey.",
            "Thoughtful analysis takes time—give yourself that gift.",
            "Your academic journey is uniquely yours—embrace it!",
            "Even experts started as beginners—keep building your expertise.",
            "Research is both an art and a science—you're developing both.",
            "The questions you're asking today will shape tomorrow's knowledge.",
            "Academic confidence grows with each challenge you overcome.",
            "Your intellectual curiosity is your greatest research tool.",
            "Great academic work balances creativity with rigor—you're on the right path.",
            "The best research papers start with genuine interest—like yours!",
            "Academic writing is a skill that improves with practice—keep going!",
            "Your research contributes to our collective understanding—that matters!",
            "Moments of confusion often precede moments of clarity. Stay with it!",
            "The joy of discovery makes the research journey worthwhile."
        ]
        encouragement = random.choice(encouragements)

        # Add encouragement, copyright text and page number
        footer_para.text = encouragement + "   |   " + FOOTER_TEXT + " | Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Style the footer text
        for run in footer_para.runs:
            run.font.size = Pt(7)  # Smaller font for footer to prevent overlap
            run.font.italic = True

        # Add page number field
        run = footer_para.add_run()
        run.font.size = Pt(7)  # Ensure consistent font size
        run.font.italic = True

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def save_document(self, filepath: str) -> str:
        """Save the document to the specified path"""
        try:
            logger.info(f"Attempting to save document to {filepath}")

            # Sanitize the filepath to remove URL-unsafe characters
            original_filepath = filepath
            directory = os.path.dirname(filepath)
            filename = os.path.basename(filepath)

            # Replace URL-unsafe characters with underscores
            sanitized_filename = self.sanitize_filename(filename)

            # Construct the sanitized filepath
            sanitized_filepath = os.path.join(directory, sanitized_filename)

            if sanitized_filepath != original_filepath:
                logger.info(f"Sanitized filepath from {original_filepath} to {sanitized_filepath}")

            # Handle case when no directory is specified (save in current directory)
            if directory:
                logger.info(f"Creating directory: {directory}")
                os.makedirs(directory, exist_ok=True)
                logger.info("Directory created successfully")

            # Save the document
            logger.info("Saving document...")
            self.document.save(sanitized_filepath)
            logger.info(f"Document saved successfully to {sanitized_filepath}")
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return False

    def sanitize_filename(self, filename: str) -> str:
        """Sanitize a filename to make it safe for saving"""
        # Replace URL-unsafe characters with underscores
        sanitized = re.sub(r'[?&+=#%:;/\\,\'\"\-]', '_', filename)
        # Replace multiple underscores with a single one
        sanitized = re.sub(r'_{2,}', '_', sanitized)
        # Limit the length of the filename
        if len(sanitized) > 100:
            base, ext = os.path.splitext(sanitized)
            sanitized = base[:96] + ext
        return sanitized

    def generate_template(self, topic: str, section: str, key_points: list, tier_name: str = None):
        """
        Generate a template prompt for section generation with tier-appropriate word counts

        Args:
            topic: The topic of the paper
            section: The section to generate
            key_points: List of key points to address
            tier_name: The subscription tier name (premium, standard, or sample)
        """
        key_points_str = "\n".join([f"- {point}" for point in key_points])

        # Get appropriate word count for this section and tier
        section_word_count = get_section_word_count(section, tier_name)
        word_count_instruction = f"Target word count for this section: {section_word_count} words."

        # Add tier-specific depth instructions
        tier_specific_guidance = ""
        if tier_name:
            if tier_name.lower() == "premium":
                tier_specific_guidance = """
                Premium Tier Requirements:
                - Provide comprehensive, in-depth analysis with substantial detail
                - Include multiple perspectives and nuanced arguments
                - Use extensive evidence and examples to support claims
                - Address potential counterarguments or limitations
                - Ensure thorough coverage of all aspects of the topic
                """
            elif tier_name.lower() == "sample":
                tier_specific_guidance = """
                Sample Tier Requirements:
                - Focus on the most essential points only
                - Be concise and direct in your writing
                - Provide a representative example of quality without depth
                - Maintain academic tone while being brief
                """

        prompt = f"""You are a PhD-level researcher and expert in {topic}. Write a rigorous, publication-quality {section} section for an academic paper. Do not use markdown formatting or special characters in headings.

{word_count_instruction}

Key points to address:
{key_points_str}

{tier_specific_guidance}

Guidelines:
1. Content Requirements:
   - Provide comprehensive coverage of the topic
   - Present clear, evidence-based arguments
   - Include relevant statistics and data when applicable
   - Address current developments in the field

2. Academic Standards:
   - Use formal academic language
   - Support claims with peer-reviewed sources
   - Follow logical flow of ideas
   - Maintain objective tone throughout

3. Citations and References:
   - Include in-text citations [Author, Year]
   - Reference recent and seminal works
   - Cite primary research studies
   - Use authoritative sources

4. Section-Specific Focus:
   For {section}:
   - Follow standard conventions for this section type
   - Provide appropriate depth and detail
   - Connect to the paper's overall narrative
   - Address key controversies or debates if relevant

5. Tables and Figures:
   - IMPORTANT: When presenting data or comparisons, use well-formatted tables with real data
   - Format tables using pipe notation (|) for clear column separation
   - Include a header row and a separator row below it
   - Example table format:

   | Variable | Value | Significance |
   |----------|-------|-------------|
   | Item 1   | 85%   | p < 0.01    |
   | Item 2   | 72%   | p < 0.05    |

Write the {section} section. Use clear headings without markdown symbols and make sure to include numbered sections and subsections. When appropriate, include properly formatted tables with relevant data."""
        return prompt

    def format_research_pack(self, research_pack: Dict) -> Document:
        """
        Format a complete research pack into a single well-organized document.
        """
        doc = Document()
        self._setup_document_styles(doc)

        # Get author name for personalization - ensure it's consistent throughout the document
        author = research_pack.get('author', '')
        user_name = author if author and author.strip() and author != "Researcher" else "Researcher"
        logger.info(f"Using author name: {user_name} for research pack")

        # Title Page
        title = research_pack["example_paper"]["title"]
        doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run("Research Pack").bold = True
        doc.add_paragraph(f"Generated: {research_pack['date_generated']}")
        # ... existing code ...
        # SECTION: Key Concepts (insert personalized, varied Lily callout after every or every other paragraph)
        if 'key_concepts' in research_pack:
            doc.add_heading("Key Concepts", 1)
            key_concepts = research_pack['key_concepts']
            callout_templates = [
                ("Lily's Brainstorm", "{user}, what new ideas or questions does the concept '{concept}' spark for you? Take a moment to jot them down!", "idea"),
                ("Lily's Writing Tip", "{user}, try explaining '{concept}' in your own words. Teaching a concept is a great way to master it!", "tip"),
                ("Lily's Confidence Booster", "{user}, don't worry if '{concept}' feels challenging at first. Every expert started as a beginner!", "confidence"),
                ("Lily's Question", "{user}, how does '{concept}' connect to other topics you've studied? Making connections deepens understanding.", "question"),
                ("Lily's Insight", "{user}, understanding '{concept}' is key to building a strong foundation for your research.", "insight"),
                ("Lily's Caution Flag", "{user}, be careful not to confuse '{concept}' with similar terms. Double-check definitions if you're unsure!", "warning"),
            ]
            for i, concept in enumerate(key_concepts):
                para = doc.add_paragraph(concept)
                para.style = 'Normal'
                # Insert a callout after every or every other paragraph (randomized)
                if i % 2 == 0 or random.choice([True, False]):
                    title, template, callout_type = random.choice(callout_templates)
                    callout_text = template.format(user=user_name, concept=concept)
                    self.add_lily_callout_box(title, callout_text, callout_type=callout_type)
        # ... existing code ...

    def add_lily_insight_with_avatar(self, insight_text, insight_type="standard"):
        """
        Add a Lily insight with a character avatar image and more personality

        This creates a more visually distinct callout box with an avatar-like presentation
        for Lily's insights.
        """
        # Create a 2-column table for avatar and content
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Make borders more subtle - using the correct approach for python-docx
        # Instead of trying to manipulate XML directly, use the table style
        table.style = 'Light Grid'  # Use a built-in light table style

        # No background shading for modern style

        # Set column widths - narrow for avatar, wider for content
        table.columns[0].width = Inches(0.8)  # Avatar column
        table.columns[1].width = self.document.sections[0].page_width - Inches(2.0)  # Content column

        # Avatar cell (first column)
        avatar_cell = table.cell(0, 0)

        # Add "Lily" text without emoji for modern style
        avatar_para = avatar_cell.paragraphs[0]
        avatar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        avatar_run = avatar_para.add_run("Lily")
        avatar_run.bold = True
        avatar_run.font.size = Pt(18)

        # Content cell (second column)
        content_cell = table.cell(0, 1)

        # Add the insight text with a header
        insight_para = content_cell.paragraphs[0]

        # Add a catchy header with personality
        header_text = "Lily's Insight"
        if insight_type == "citation":
            header_text = "Citation Tip from Lily"
        elif insight_type == "idea":
            header_text = "Lily's Brainstorm"
        elif insight_type == "counterpoint":
            header_text = "Consider This..."
        elif insight_type == "confidence":
            header_text = "Lily's Confidence Booster"
        elif insight_type == "warning":
            header_text = "Lily's Caution Flag"
        elif insight_type == "tip":
            header_text = "Lily's Writing Tip"
        elif insight_type == "question":
            header_text = "Question to Consider"

        # Add a header with modern styling
        header_run = insight_para.add_run(f"{header_text}:\n")
        header_run.bold = True
        header_run.font.size = Pt(12)

        # Add the actual insight text
        text_run = insight_para.add_run(insight_text)
        text_run.italic = True
        text_run.font.size = Pt(11)

        # Add spacing paragraph after the table
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_lily_insight(self, insight_text, insight_type="standard"):
        """Add a Lily insight with modern styling (no emoji)"""
        # Use the new avatar version instead for better styling
        return self.add_lily_insight_with_avatar(insight_text, insight_type)

    def add_research_highlight(self, research_type, source, content, year="2023"):
        """
        Add a research highlight box with citation information.

        Args:
            research_type: Type of research (e.g., "Research Methodology")
            source: Source of the research
            content: The highlight content
            year: Publication year
        """
        # Create a table for the highlight with borders
        table = self.document.add_table(rows=2, cols=1)
        table.style = 'Table Grid'

        # Get the header cell
        header_cell = table.cell(0, 0)

        # Add light yellow shading to the header cell background
        shading_element = header_cell._element.tcPr.xpath('./w:shd')
        if not shading_element:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFE0"/>')
            header_cell._element.tcPr.append(shading_element)
        else:
            shading_element[0].set(qn('w:fill'), 'FFFFE0')  # Light yellow background

        # Add header content
        header_para = header_cell.paragraphs[0]
        header_run = header_para.add_run(f"Research Highlight: {research_type}")
        header_run.bold = True
        header_run.font.size = Pt(12)

        # Get the content cell
        content_cell = table.cell(1, 0)

        # Add content
        content_para = content_cell.paragraphs[0]
        content_para.text = content
        content_para.add_run(f"\n\nSource: {source}, {year}")

        # Add spacing after the highlight
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_learning_objectives_box(self, topic, objectives):
        """
        Add a 'You Will Learn...' box at the top of the Overview section.

        Args:
            topic: The research topic
            objectives: List of learning objectives
        """
        # Create a table for the box with borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add title
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(f"Quick-Glance: What You Will Learn")
        title_run.bold = True
        title_run.font.size = Pt(14)

        # Add subtitle
        subtitle_para = cell.add_paragraph()
        subtitle_run = subtitle_para.add_run(f"This research pack will help you develop a comprehensive understanding of {topic}.")
        subtitle_run.italic = True

        # Add objectives as bullet points
        for objective in objectives:
            bullet_para = cell.add_paragraph(style='List Bullet')
            bullet_para.text = objective
            bullet_para.paragraph_format.left_indent = Inches(0.25)

        # Add spacing after the box
        self.document.add_paragraph()

        return table

    def add_info_box(self, title, content, box_type="info"):
        """
        Add a visually distinct information box with a title and content.

        Args:
            title: The title of the info box
            content: The main text content
            box_type: Type of box (info, warning, tip, research, fact)
        """
        # Determine box styling based on type
        box_style_name = f'{box_type.capitalize()} Box'
        if box_style_name not in self.document.styles:
            box_style = self.document.styles.add_style(box_style_name, WD_STYLE_TYPE.PARAGRAPH)
            box_style.font.size = Pt(11)
            box_style.paragraph_format.space_after = Pt(6)
            box_style.paragraph_format.space_before = Pt(6)
            box_style.paragraph_format.left_indent = Inches(0.25)
            box_style.paragraph_format.right_indent = Inches(0.25)

        # Create title style if needed
        title_style_name = f'{box_type.capitalize()} Box Title'
        if title_style_name not in self.document.styles:
            title_style = self.document.styles.add_style(title_style_name, WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(12)
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(6)
            title_style.paragraph_format.left_indent = Inches(0.25)

        # Create a table for the box with borders - using Light Grid for more subtle borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell and set its properties
        cell = table.cell(0, 0)

        # Add title without icon for modern style
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_para.style = title_style_name

        # Add content
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = box_style_name

        # Add spacing after the box
        self.document.add_paragraph()

        return table

    def add_lily_callout(self, content):
        """
        Add a simple Lily callout with the given content.
        This is used in the original research pack orchestrator.

        Args:
            content: The callout content
        """
        # Create a table for the callout with borders - using Light Grid for more subtle borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # No background shading for modern style

        # Add content with clean styling (no emoji)
        content_para = cell.paragraphs[0]

        # Add title
        title_run = content_para.add_run("Lily's Insight:\n")
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add content
        content_run = content_para.add_run(content)
        content_run.italic = True
        content_run.font.size = Pt(11)

        # Add spacing after the callout
        self.document.add_paragraph().paragraph_format.space_after = Pt(12)

        return table

    def add_lily_callout_box(self, title, content, callout_type="guidance", source=None):
        """
        Add a special Lily callout box with visual flair

        Args:
            title: The title of the callout
            content: The main text content
            callout_type: Type of callout (tip, insight, question, warning, confidence, brainstorm, research, guidance, reflection, connection, example, definition, caution)
            source: Source of the callout (None for hardcoded, 'llm_generated' for LLM-generated)
        """
        # Create a table for the callout with borders - using Light Grid for more subtle borders
        table = self.document.add_table(rows=1, cols=1)

        # Use different style for LLM-generated callouts to make them visually distinct
        if source == "llm_generated":
            table.style = 'Light List Accent 1'  # Use a different style for LLM-generated callouts
        else:
            table.style = 'Light Grid'  # Default style for hardcoded callouts

        # Get the cell
        cell = table.cell(0, 0)

        # Add title without icon for modern style
        title_para = cell.paragraphs[0]
        title_para.text = ""

        # Different title format for LLM-generated callouts
        if source == "llm_generated":
            title_run = title_para.add_run(f"{title}")
        else:
            title_run = title_para.add_run(f"Lily's {callout_type.capitalize()}: {title}")

        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add the content
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = 'Normal'

        # Add spacing after the callout
        self.document.add_paragraph()

        return table

    def add_research_highlight(self, research_title, source, content, year=None):
        """
        Add a research highlight box that showcases key research findings

        Args:
            research_title: Title of the research paper
            source: Source of the research (authors/journal)
            content: Key findings or content to highlight
            year: Publication year (optional)
        """
        # Create a table for the research highlight with borders - using Light Grid for more subtle borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add research title without icon for modern style
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(f"Research: {research_title}")
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add source information
        source_text = source
        if year:
            source_text += f" ({year})"

        source_para = cell.add_paragraph()
        source_run = source_para.add_run(source_text)
        source_run.italic = True
        source_run.font.size = Pt(10)

        # Add the key findings
        content_para = cell.add_paragraph()
        content_para.text = content
        content_para.style = 'Normal'

        # Add spacing after the research highlight
        self.document.add_paragraph()

        return table

    def add_side_note(self, text, note_type="tip"):
        """
        Add a side note with different styling than the main content

        Args:
            text: The side note text
            note_type: Type of note (tip, note, important)
        """
        # Create a paragraph for the side note
        paragraph = self.document.add_paragraph()

        # Add text without icon for modern style
        run = paragraph.add_run(text)
        run.italic = True

        # Style the paragraph
        if f'{note_type.capitalize()} Note' not in self.document.styles:
            note_style = self.document.styles.add_style(f'{note_type.capitalize()} Note', WD_STYLE_TYPE.PARAGRAPH)
            note_style.font.size = Pt(10)
            note_style.font.italic = True
            note_style.paragraph_format.left_indent = Inches(0.5)
            note_style.paragraph_format.right_indent = Inches(0.5)
            note_style.paragraph_format.space_before = Pt(6)
            note_style.paragraph_format.space_after = Pt(6)

        paragraph.style = f'{note_type.capitalize()} Note'

        return paragraph

    def add_visual_separator(self, separator_type="stars"):
        """
        Add a visual separator to break up content sections

        Args:
            separator_type: Type of separator (stars, line, dots)
        """
        paragraph = self.document.add_paragraph()

        # Set the separator text based on type
        if separator_type == "stars":
            separator_text = "* * * * *"
        elif separator_type == "line":
            separator_text = "―――――――――――――――――――"
        elif separator_type == "dots":
            separator_text = "• • • • •"
        else:
            separator_text = "―――――――――――――――――――"

        # Add the separator with center alignment
        run = paragraph.add_run(separator_text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing before and after
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        return paragraph

    def add_research_data_section(self, title, data_points, source=None):
        """
        Add a section with research data presented in a visually appealing way

        Args:
            title: Title for the research data section
            data_points: List of research findings or data points
            source: Data source (optional)
        """
        # Create a table for the data section - using Light Grid for more subtle borders
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add title without icon for modern style
        title_para = cell.paragraphs[0]
        title_para.text = ""
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add source if provided
        if source:
            source_para = cell.add_paragraph()
            source_run = source_para.add_run(f"Source: {source}")
            source_run.italic = True
            source_run.font.size = Pt(9)

        # Add each data point as bullet point
        for data_point in data_points:
            point_para = cell.add_paragraph()
            point_para.style = 'List Bullet'
            point_para.text = data_point

        # Add spacing after the data section
        self.document.add_paragraph()

        return table

    def add_research_plan_table(self, topic, education_level):
        """
        Add a 4-Week Research Plan table to the document

        Args:
            topic: The research topic
            education_level: The education level (university, postgraduate, etc.)
        """
        # Add heading
        self.add_heading("4-Week Research Plan", level=2)

        # Add description
        self.add_paragraph(f"Use this customized research plan to organize your work on '{topic}'. Adjust the timeline as needed to fit your schedule.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=5, cols=3)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.0)  # Week column
        table.columns[1].width = Inches(2.5)  # Tasks column
        table.columns[2].width = Inches(2.5)  # Outcomes column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Week"
        header_cells[1].text = "Tasks"
        header_cells[2].text = "Expected Outcomes"

        # Format headers
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content for each week
        weeks = [
            ("Week 1",
             f"• Explore {topic} broadly\n• Identify 3-5 subtopics\n• Gather initial sources\n• Create research questions",
             "• List of potential research questions\n• Initial bibliography (5-7 sources)\n• Defined scope of research"),
            ("Week 2",
             f"• Read and take notes on sources\n• Refine research question\n• Identify key debates in {topic}\n• Create detailed outline",
             "• Finalized research question\n• Annotated bibliography\n• Working outline with main points"),
            ("Week 3",
             f"• Draft introduction and literature review\n• Analyze key arguments\n• Gather additional sources as needed\n• Organize evidence",
             "• Complete first draft (intro and lit review)\n• Organized notes for main body\n• Clear argument structure"),
            ("Week 4",
             f"• Complete full draft\n• Revise for clarity and coherence\n• Check citations and formatting\n• Proofread final document",
             "• Complete final paper\n• Properly formatted citations\n• Polished, error-free document")
        ]

        # Fill in the table
        for i, (week, tasks, outcomes) in enumerate(weeks):
            row = table.rows[i+1]
            row.cells[0].text = week
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[0].paragraphs[0].runs[0].bold = True

            # Add tasks with bullet points
            tasks_para = row.cells[1].paragraphs[0]
            tasks_para.text = ""
            for task in tasks.split("\n"):
                if tasks_para.text:
                    tasks_para = row.cells[1].add_paragraph()
                tasks_para.text = task

            # Add outcomes with bullet points
            outcomes_para = row.cells[2].paragraphs[0]
            outcomes_para.text = ""
            for outcome in outcomes.split("\n"):
                if outcomes_para.text:
                    outcomes_para = row.cells[2].add_paragraph()
                outcomes_para.text = outcome

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_topic_mind_map_table(self, topic):
        """
        Add a Topic Mind Map table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Topic Mind Map", level=2)

        # Add description
        self.add_paragraph(f"Use this mind map to explore different aspects of '{topic}'. Fill in the empty cells with your ideas and connections.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=5, cols=3)
        table.style = 'Light Grid'

        # Merge the center cell for the main topic
        center_cell = table.cell(2, 1)
        center_cell.text = topic
        center_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_cell.paragraphs[0].runs[0].bold = True

        # Add subtopic placeholders
        subtopics = [
            (1, 0, "Historical Context"),
            (1, 2, "Current Debates"),
            (0, 1, "Key Theories"),
            (3, 0, "Methodological Approaches"),
            (3, 2, "Practical Applications"),
            (4, 1, "Future Directions")
        ]

        for row, col, text in subtopics:
            cell = table.cell(row, col)
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add empty cells for user to fill in
        empty_cells = [
            (0, 0, "[Your notes]"),
            (0, 2, "[Your notes]"),
            (2, 0, "[Your notes]"),
            (2, 2, "[Your notes]"),
            (4, 0, "[Your notes]"),
            (4, 2, "[Your notes]")
        ]

        for row, col, text in empty_cells:
            cell = table.cell(row, col)
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_thesis_builder_table(self, topic):
        """
        Add a Thesis/Argument Builder table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Thesis/Argument Builder", level=2)

        # Add description
        self.add_paragraph(f"Use this template to develop a strong thesis statement and supporting arguments for your research on '{topic}'.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=7, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Component column
        table.columns[1].width = Inches(4.0)  # Content column

        # Add rows with content
        rows_content = [
            ("Research Question", f"What specific aspect of {topic} do you want to investigate?"),
            ("Main Claim/Thesis", "Your central argument in one clear sentence:"),
            ("Supporting Argument 1", "First key point that supports your thesis:"),
            ("Supporting Argument 2", "Second key point that supports your thesis:"),
            ("Supporting Argument 3", "Third key point that supports your thesis:"),
            ("Counterargument", "What opposing viewpoint will you address?"),
            ("Conclusion", "How does your argument contribute to understanding of the topic?")
        ]

        # Fill in the table
        for i, (component, content) in enumerate(rows_content):
            row = table.rows[i]
            row.cells[0].text = component
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = content
            row.cells[1].paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_literature_review_matrix(self, topic):
        """
        Add a Literature Review Matrix table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Literature Review Matrix", level=2)

        # Add description
        self.add_paragraph(f"Use this matrix to organize key sources for your research on '{topic}'. Add rows as needed for additional sources.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=6, cols=5)
        table.style = 'Light Grid'

        # Add headers
        header_cells = table.rows[0].cells
        headers = ["Source Citation", "Main Arguments", "Methodology", "Key Findings", "Relevance to Your Research"]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add empty rows for user to fill in
        for row in range(1, 6):
            for col in range(5):
                table.cell(row, col).text = "[Enter your information]" if col == 0 else ""

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_argument_planning_grid(self, topic):
        """
        Add an Argument Planning Grid table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Argument Planning Grid", level=2)

        # Add description
        self.add_paragraph(f"Use this grid to map out your arguments, evidence, counterarguments, and rebuttals for your research on '{topic}'.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=4, cols=4)
        table.style = 'Light Grid'

        # Add headers
        header_cells = table.rows[0].cells
        headers = ["Main Argument", "Supporting Evidence", "Counterargument", "Rebuttal"]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add empty rows for user to fill in
        for row in range(1, 4):
            for col in range(4):
                table.cell(row, col).text = ""

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_editable_timeline(self, topic):
        """
        Add an Editable Timeline table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Editable Research Timeline", level=2)

        # Add description
        self.add_paragraph(f"Customize this timeline for your research project on '{topic}'. Add specific dates and milestones based on your schedule.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=7, cols=3)
        table.style = 'Light Grid'

        # Add headers
        header_cells = table.rows[0].cells
        headers = ["Date/Week", "Research Phase", "Tasks & Milestones"]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add phases with empty date cells
        phases = [
            "Exploration Phase",
            "Literature Review",
            "Data Collection/Analysis",
            "Drafting",
            "Revision",
            "Final Submission"
        ]

        for i, phase in enumerate(phases):
            row = table.rows[i+1]
            row.cells[0].text = "[Enter date]"
            row.cells[0].paragraphs[0].runs[0].italic = True
            row.cells[1].text = phase
            row.cells[2].text = "[Enter your tasks]"
            row.cells[2].paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_emergency_research_plan(self):
        """
        Add an Emergency Research Recovery Plan table to the document
        """
        # Add heading
        self.add_heading("Emergency Research Recovery Plan", level=2)

        # Add description
        self.add_paragraph("Use this checklist when you encounter obstacles in your research process. These strategies can help you get back on track.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Problem column
        table.columns[1].width = Inches(4.0)  # Solution column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Common Problem"
        header_cells[1].text = "Recovery Strategy"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content rows
        rows_content = [
            ("Writer's Block", "• Break task into smaller chunks\n• Change your environment\n• Free-write for 15 minutes without editing\n• Discuss your ideas with someone else"),
            ("Overwhelmed by Sources", "• Create categories for your sources\n• Focus on most recent/relevant first\n• Use the Literature Review Matrix to organize\n• Limit yourself to 3-5 sources per section"),
            ("Unclear Argument", "• Return to your Thesis Builder\n• Write a one-paragraph summary of your argument\n• Test your thesis with a skeptical friend\n• Create a mind map of your key points"),
            ("Time Management Issues", "• Reset your timeline with realistic deadlines\n• Use the Pomodoro technique (25 min work, 5 min break)\n• Identify and eliminate distractions\n• Schedule specific research tasks"),
            ("Feedback Setback", "• Separate criticism from your self-worth\n• Categorize feedback as major/minor\n• Create an action plan for revisions\n• Seek clarification on unclear feedback")
        ]

        # Fill in the table
        for i, (problem, solution) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = problem
            row.cells[0].paragraphs[0].runs[0].bold = True

            # Add solutions with bullet points
            solutions_para = row.cells[1].paragraphs[0]
            solutions_para.text = ""
            for sol in solution.split("\n"):
                if solutions_para.text:
                    solutions_para = row.cells[1].add_paragraph()
                solutions_para.text = sol

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_reflection_prompts(self, topic):
        """
        Add Reflection & Self-Assessment Prompts table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Reflection & Self-Assessment Prompts", level=2)

        # Add description
        self.add_paragraph(f"Use these prompts to reflect on your research process and assess your work on '{topic}'. Regular reflection improves your critical thinking and research skills.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Stage column
        table.columns[1].width = Inches(4.0)  # Prompts column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Research Stage"
        header_cells[1].text = "Reflection Prompts"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content rows
        rows_content = [
            ("Topic Selection", "• Why did I choose this topic?\n• What aspects of this topic interest me most?\n• What do I hope to discover through my research?\n• How does this topic connect to my broader academic goals?"),
            ("Literature Review", "• What patterns am I noticing across sources?\n• Which sources have been most valuable and why?\n• What gaps or contradictions exist in the literature?\n• How has my understanding of the topic evolved?"),
            ("Argument Development", "• How clear and focused is my thesis statement?\n• How well do my supporting points connect to my thesis?\n• What counterarguments have I considered?\n• What is the strongest and weakest part of my argument?"),
            ("Writing Process", "• What writing strategies have been most effective?\n• Where am I getting stuck in the writing process?\n• How well am I integrating evidence into my writing?\n• What feedback has been most helpful?"),
            ("Final Assessment", "• How has my thinking on this topic changed?\n• What are the strongest aspects of my research?\n• What would I do differently next time?\n• What new questions has this research raised?")
        ]

        # Fill in the table
        for i, (stage, prompts) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = stage
            row.cells[0].paragraphs[0].runs[0].bold = True

            # Add prompts with bullet points
            prompts_para = row.cells[1].paragraphs[0]
            prompts_para.text = ""
            for prompt in prompts.split("\n"):
                if prompts_para.text:
                    prompts_para = row.cells[1].add_paragraph()
                prompts_para.text = prompt

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_glossary_builder(self, topic):
        """
        Add a Glossary Builder table to the document

        Args:
            topic: The research topic
        """
        # Add heading
        self.add_heading("Glossary Builder", level=2)

        # Add description
        self.add_paragraph(f"Use this template to create a glossary of key terms related to '{topic}'. A clear understanding of terminology is essential for effective research.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=11, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Term column
        table.columns[1].width = Inches(4.0)  # Definition column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Term"
        header_cells[1].text = "Definition & Source"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add empty rows for user to fill in
        for row in range(1, 11):
            table.cell(row, 0).text = f"[Term {row}]"
            table.cell(row, 0).paragraphs[0].runs[0].italic = True
            table.cell(row, 1).text = "[Enter definition and citation]"
            table.cell(row, 1).paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_peer_review_checklist(self):
        """
        Add a Peer Review Checklist table to the document
        """
        # Add heading
        self.add_heading("Peer Review Checklist", level=2)

        # Add description
        self.add_paragraph("Use this checklist when reviewing your own work or when asking peers to review your research. Constructive feedback improves the quality of your final paper.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=6, cols=3)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Component column
        table.columns[1].width = Inches(3.0)  # Questions column
        table.columns[2].width = Inches(1.0)  # Rating column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Component"
        header_cells[1].text = "Review Questions"
        header_cells[2].text = "Rating (1-5)"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content rows
        rows_content = [
            ("Thesis & Argument", "• Is the thesis clear and specific?\n• Are the main arguments well-supported?\n• Is the argument logical and coherent?\n• Are counterarguments addressed?"),
            ("Evidence & Sources", "• Is evidence relevant and sufficient?\n• Are sources credible and appropriate?\n• Are quotes integrated effectively?\n• Are citations formatted correctly?"),
            ("Structure & Organization", "• Does the introduction effectively set up the paper?\n• Do paragraphs flow logically?\n• Are transitions between sections smooth?\n• Does the conclusion synthesize key points?"),
            ("Writing Style & Clarity", "• Is the writing clear and concise?\n• Is academic language used appropriately?\n• Are sentences varied and well-constructed?\n• Is the tone consistent and appropriate?"),
            ("Overall Assessment", "• What are the paper's greatest strengths?\n• What are the top 3 areas for improvement?\n• How effectively does the paper address the research question?\n• What additional sources or perspectives should be considered?")
        ]

        # Fill in the table
        for i, (component, questions) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = component
            row.cells[0].paragraphs[0].runs[0].bold = True

            # Add questions with bullet points
            questions_para = row.cells[1].paragraphs[0]
            questions_para.text = ""
            for question in questions.split("\n"):
                if questions_para.text:
                    questions_para = row.cells[1].add_paragraph()
                questions_para.text = question

            row.cells[2].text = "____"
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_source_tracking_table(self):
        """
        Add a Source Tracking Table to the document
        """
        # Add heading
        self.add_heading("Source Tracking Table", level=2)

        # Add description
        self.add_paragraph("Use this table to keep track of your sources and how you plan to use them in your research.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=11, cols=4)
        table.style = 'Light Grid'

        # Add headers
        header_cells = table.rows[0].cells
        headers = ["Source Citation", "Key Points/Findings", "Relevance to Research", "Page Numbers/Notes"]

        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add empty rows for user to fill in
        for row in range(1, 11):
            for col in range(4):
                placeholder = "[Enter citation]" if col == 0 else ""
                table.cell(row, col).text = placeholder
                if placeholder:
                    table.cell(row, col).paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_introduction_builder(self):
        """
        Add an Introduction Builder table to the document
        """
        # Add heading
        self.add_heading("Introduction Builder", level=2)

        # Add description
        self.add_paragraph("Use this template to craft a strong introduction for your research paper. A well-structured introduction establishes context, presents your thesis, and outlines your approach.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Component column
        table.columns[1].width = Inches(4.0)  # Content column

        # Add rows with content
        rows_content = [
            ("Hook/Opening", "Start with an engaging statement that captures attention:"),
            ("Context/Background", "Provide necessary background information on your topic:"),
            ("Research Gap", "Identify the gap or problem your research addresses:"),
            ("Research Question", "State your primary research question(s):"),
            ("Thesis Statement", "Present your main argument or position:"),
            ("Paper Structure", "Briefly outline how your paper will develop:")
        ]

        # Fill in the table
        for i, (component, content) in enumerate(rows_content):
            row = table.rows[i]
            row.cells[0].text = component
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = content
            row.cells[1].paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_paragraph_planner(self):
        """
        Add a Paragraph Planner (PEEL Format) table to the document
        """
        # Add heading
        self.add_heading("Paragraph Planner (PEEL Format)", level=2)

        # Add description
        self.add_paragraph("Use this template to plan your paragraphs using the PEEL format (Point, Evidence, Explanation, Link). This structure helps create clear, focused paragraphs that advance your argument.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Component column
        table.columns[1].width = Inches(4.5)  # Content column

        # Add rows with content
        rows_content = [
            ("P - Point", "State your main point or claim for this paragraph:"),
            ("E - Evidence", "Provide supporting evidence (quotes, data, examples):"),
            ("E - Explanation", "Explain how your evidence supports your point:"),
            ("L - Link", "Connect back to your thesis or to the next paragraph:")
        ]

        # Fill in the table
        for i, (component, content) in enumerate(rows_content):
            row = table.rows[i]
            row.cells[0].text = component
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = content
            row.cells[1].paragraphs[0].runs[0].italic = True

        # Add note about using multiple copies
        note_cell = table.cell(4, 0)
        note_cell.merge(table.cell(4, 1))
        note_cell.text = "Note: Make copies of this template for each main paragraph in your paper. Aim for one clear point per paragraph."
        note_cell.paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_thesis_alignment_check(self):
        """
        Add a Thesis Alignment Check table to the document
        """
        # Add heading
        self.add_heading("Thesis Alignment Check", level=2)

        # Add description
        self.add_paragraph("Use this tool to ensure all parts of your paper align with and support your thesis statement. This helps maintain focus and coherence throughout your research.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=7, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(2.0)  # Component column
        table.columns[1].width = Inches(4.0)  # Content column

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Component"
        header_cells[1].text = "Alignment with Thesis"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add rows with content
        rows_content = [
            ("Thesis Statement", "Write your complete thesis statement here:"),
            ("Introduction", "How does your introduction set up and lead to your thesis?"),
            ("Body Paragraph 1", "How does this paragraph support your thesis?"),
            ("Body Paragraph 2", "How does this paragraph support your thesis?"),
            ("Body Paragraph 3", "How does this paragraph support your thesis?"),
            ("Conclusion", "How does your conclusion reinforce your thesis?")
        ]

        # Fill in the table
        for i, (component, content) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = component
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = content
            row.cells[1].paragraphs[0].runs[0].italic = True

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_citation_cheat_sheet(self):
        """
        Add a Citation Cheat Sheet table to the document
        """
        # Add heading
        self.add_heading("Citation Cheat Sheet", level=2)

        # Add description
        self.add_paragraph("Use this quick reference guide for formatting citations in your research paper. Always check your specific style guide for complete requirements.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=4, cols=3)
        table.style = 'Light Grid'

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Source Type"
        header_cells[1].text = "APA Format"
        header_cells[2].text = "MLA Format"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content rows
        rows_content = [
            ("Book", "Author, A. A. (Year). Title of book. Publisher.", "Author, First Name. Title of Book. Publisher, Year.")
            ,
            ("Journal Article", "Author, A. A. (Year). Title of article. Journal Name, Volume(Issue), page range. DOI", "Author, First Name. 'Title of Article.' Journal Name, vol. number, no. issue, Year, pp. page range.")
            ,
            ("Website", "Author, A. A. (Year, Month Day). Title of page. Site Name. URL", "Author, First Name. 'Title of Page.' Site Name, Day Month Year, URL.")
        ]

        # Fill in the table
        for i, (source_type, apa, mla) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = source_type
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = apa
            row.cells[2].text = mla

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def add_one_week_crash_plan(self):
        """
        Add a One-Week Crash Plan table to the document
        """
        # Add heading
        self.add_heading("One-Week Crash Plan (Panic Mode)", level=2)

        # Add description
        self.add_paragraph("If you're running out of time, use this emergency plan to complete your research paper in one week. While not ideal, this structured approach will help you produce the best possible paper under time constraints.")

        # Create the table with Light Grid style for more subtle borders
        table = self.document.add_table(rows=8, cols=2)
        table.style = 'Light Grid'

        # Set column widths
        table.columns[0].width = Inches(1.5)  # Day column
        table.columns[1].width = Inches(4.5)  # Tasks column

        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Day"
        header_cells[1].text = "Emergency Tasks"

        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True

        # Add content rows
        rows_content = [
            ("Day 1", "• Narrow your topic immediately\n• Formulate a specific research question\n• Find 5-7 key sources (focus on recent academic sources)\n• Create a basic outline with thesis statement")
            ,
            ("Day 2", "• Skim all sources and take focused notes\n• Identify key quotes and evidence\n• Refine your thesis and outline\n• Draft introduction")
            ,
            ("Day 3", "• Write first body section\n• Begin second body section\n• Create bibliography entries as you go")
            ,
            ("Day 4", "• Complete all body sections\n• Ensure each paragraph has a clear point\n• Add transitions between sections")
            ,
            ("Day 5", "• Write conclusion\n• Review argument flow and logic\n• Add any missing citations")
            ,
            ("Day 6", "• Complete full revision for content and clarity\n• Check all citations match reference list\n• Format document according to requirements")
            ,
            ("Day 7", "• Final proofreading for grammar and spelling\n• Check formatting requirements\n• Submit paper")
        ]

        # Fill in the table
        for i, (day, tasks) in enumerate(rows_content):
            row = table.rows[i+1]
            row.cells[0].text = day
            row.cells[0].paragraphs[0].runs[0].bold = True

            # Add tasks with bullet points
            tasks_para = row.cells[1].paragraphs[0]
            tasks_para.text = ""
            for task in tasks.split("\\n"):
                if tasks_para.text:
                    tasks_para = row.cells[1].add_paragraph()
                tasks_para.text = task

        # Add spacing after the table
        self.document.add_paragraph()

        return table

    def generate_mind_map(self, topic):
        """
        Generate a mind map for the given topic

        Args:
            topic: The research topic

        Returns:
            Path to the generated mind map image
        """
        try:
            # Import the diagram orchestrator directly
            import os
            import sys
            import traceback

            # Add the project root to the path to find diagram_orchestrator
            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../../'))
            logger.info(f"Project root path: {project_root}")

            if project_root not in sys.path:
                sys.path.append(project_root)
                logger.info(f"Added project root to sys.path: {sys.path}")

            # Use the diagram orchestrator's mind map generator
            try:
                # Import directly from diagram_orchestrator
                from diagram_orchestrator.orchestrator import generate_mind_map
                logger.info(f"Successfully imported diagram_orchestrator.orchestrator module")
                logger.info(f"Using diagram orchestrator mind map generator for topic: {topic}")

                # Generate the mind map
                result = generate_mind_map(topic)
                logger.info(f"Mind map generation result: {result}")
                return result
            except ImportError as ie:
                logger.error(f"Import error: {str(ie)}")
                logger.error(traceback.format_exc())
                logger.error("CRITICAL ERROR: Could not import diagram_orchestrator.orchestrator")
                logger.error("This is a critical error - mind maps will not work without this module")
                return None
        except Exception as e:
            logger.error(f"Error generating mind map: {str(e)}")
            logger.error(traceback.format_exc())
            return None
    def generate_research_process_diagram(self, topic):
        """
        Generate a research process diagram for the given topic

        Args:
            topic: The research topic

        Returns:
            Path to the generated diagram image
        """
        try:
            # Import the diagram orchestrator directly
            import os
            import sys
            import traceback

            # Add the project root to the path to find diagram_orchestrator
            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../../'))
            logger.info(f"Project root path: {project_root}")

            if project_root not in sys.path:
                sys.path.append(project_root)
                logger.info(f"Added project root to sys.path: {sys.path}")

            # Use the diagram orchestrator's question breakdown generator
            try:
                # Import directly from diagram_orchestrator
                from diagram_orchestrator.orchestrator import generate_research_process_diagram
                logger.info(f"QUESTION BREAKDOWN DEBUG: Successfully imported diagram_orchestrator.orchestrator module")
                logger.info(f"QUESTION BREAKDOWN DEBUG: Using diagram orchestrator question breakdown generator for topic: {topic}")

                # Generate the question breakdown diagram
                result = generate_research_process_diagram(topic)
                logger.info(f"QUESTION BREAKDOWN DEBUG: Question breakdown generation result: {result}")
                return result
            except ImportError as ie:
                logger.error(f"Import error: {str(ie)}")
                logger.error(traceback.format_exc())
                logger.error("CRITICAL ERROR: Could not import diagram_orchestrator.orchestrator")
                logger.error("This is a critical error - question breakdown diagrams will not work without this module")

                # Fall back to the legacy implementation
                try:
                    logger.warning("Falling back to legacy research process diagram generator")
                    current_dir = os.path.dirname(os.path.abspath(__file__))
                    if current_dir not in sys.path:
                        sys.path.append(current_dir)
                    from fix_diagrams import generate_research_process_diagram as legacy_generator
                    return legacy_generator(topic)
                except Exception as fallback_error:
                    logger.error(f"Error in fallback research process diagram generator: {str(fallback_error)}")
                    return None
        except Exception as e:
            logger.error(f"Error generating research process diagram: {str(e)}")
            logger.error(traceback.format_exc())
            return None

    def add_diagram_to_document(self, diagram_path: str, caption: str = None, use_landscape: bool = False):
        """
        Add a diagram to the document

        Args:
            diagram_path: Path to the diagram image file
            caption: Optional caption for the diagram
            use_landscape: Whether to use landscape orientation for the diagram

        Returns:
            True if the diagram was added successfully, False otherwise
        """
        try:
            logger.info(f"DIAGRAM DEBUG: add_diagram_to_document called with diagram_path={diagram_path}, caption={caption}, use_landscape={use_landscape}")

            if not diagram_path:
                logger.error(f"DIAGRAM DEBUG: diagram_path is None or empty")
                return False

            if not os.path.exists(diagram_path):
                logger.error(f"DIAGRAM DEBUG: diagram_path does not exist: {diagram_path}")
                return False

            if diagram_path and os.path.exists(diagram_path):
                logger.info(f"DIAGRAM DEBUG: diagram file exists at {diagram_path}")

                # Check if this is a diagram that should use landscape orientation (based on filename)
                diagram_filename = os.path.basename(diagram_path).lower()
                is_mind_map = 'mind_map' in diagram_filename
                is_journey_map = 'journey_map' in diagram_filename or 'leo_journey_map' in diagram_filename
                is_process_diagram = 'process' in diagram_filename and 'diagram' in diagram_filename

                # Use landscape orientation for certain diagrams by default
                if is_mind_map or is_journey_map or is_process_diagram:
                    use_landscape = True
                    logger.info(f"DIAGRAM DEBUG: Using landscape orientation for diagram: {diagram_filename}")
                    logger.info(f"DIAGRAM DEBUG: is_mind_map={is_mind_map}, is_journey_map={is_journey_map}, is_process_diagram={is_process_diagram}")

                # Use landscape formatter if requested
                if use_landscape:
                    try:
                        # Import the landscape formatter
                        import sys
                        sys.path.append('/home/admin/projects/kdd')
                        from diagram_orchestrator.utils.landscape_formatter import add_landscape_diagram_page

                        # Add figure number
                        self.figure_count = getattr(self, 'figure_count', 0) + 1

                        # Update caption with figure number if provided
                        if caption:
                            # Clean up caption
                            clean_caption = caption.replace('Figure:', '').strip()  # Remove any existing 'Figure:' prefix
                            clean_caption = clean_caption.replace('##', '').strip()  # Remove markdown heading syntax
                            full_caption = f'Figure {self.figure_count}. {self.clean_markdown(clean_caption)}'
                        else:
                            full_caption = None

                        # Add the diagram in landscape orientation
                        result = add_landscape_diagram_page(self.document, diagram_path, full_caption)
                        if result:
                            logger.info(f"Successfully added diagram in landscape orientation: {diagram_path}")
                            return True
                        else:
                            logger.warning(f"Failed to add diagram in landscape orientation, falling back to portrait: {diagram_path}")
                    except Exception as landscape_error:
                        logger.error(f"Error using landscape formatter: {str(landscape_error)}")
                        logger.error("LANDSCAPE FORMATTER FAILED - This is a critical error for mind maps")
                        # Do not fall back to portrait orientation for mind maps
                        if is_mind_map:
                            logger.error(f"Mind map requires landscape orientation, cannot continue with portrait fallback")
                            return False
                        else:
                            logger.info("Falling back to portrait orientation for non-mind map diagram")

                # Standard portrait orientation (only used for non-mind map diagrams or if use_landscape is False)
                # Add spacing before diagram
                self.document.add_paragraph().paragraph_format.space_before = Pt(12)

                # Add the diagram image
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                picture = run.add_picture(diagram_path, width=Inches(6))

                # Set width while maintaining aspect ratio
                target_width = Inches(6.0)  # Standard width for diagrams
                if hasattr(picture, 'height') and hasattr(picture, 'width') and picture.width > 0:
                    aspect_ratio = picture.height / picture.width
                    picture.width = target_width
                    picture.height = int(target_width * aspect_ratio)

                # Add figure number
                self.figure_count = getattr(self, 'figure_count', 0) + 1

                # Add caption if provided
                if caption:
                    # Clean up caption and add with automatic numbering
                    clean_caption = caption.replace('Figure:', '').strip()  # Remove any existing 'Figure:' prefix
                    clean_caption = clean_caption.replace('##', '').strip()  # Remove markdown heading syntax
                    caption_para = self.document.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.style = 'Caption'
                    caption_para.add_run(f'Figure {self.figure_count}. {self.clean_markdown(clean_caption)}')

                # Add spacing after diagram
                self.document.add_paragraph().paragraph_format.space_after = Pt(12)

                logger.info(f"Successfully added diagram from {diagram_path}")
                return True
            else:
                logger.warning(f"Diagram file not found: {diagram_path}")
                return False
        except Exception as e:
            logger.error(f"Error adding diagram to document: {str(e)}")
            return False

    def format_paper(self, paper, author="Researcher", institution="University", education_level="university", user_id=None, paper_id=None):
        """
        Format a paper into a professional document.

        Args:
            paper: The paper object containing title, sections, etc.
            author: The author's name
            institution: The institution name
            education_level: The education level (university, postgraduate, etc.)
            user_id: The ID of the user who owns the paper
            paper_id: The ID of the paper

        Returns:
            Dictionary with paths to the saved documents (docx and pdf)
        """
        try:
            # Get the title and sections from the paper
            title = paper.get("title", "Research Paper")
            sections = paper.get("sections", {})

            # Create the title page
            current_date = datetime.now().strftime("%B %d, %Y")
            self.create_title_page(title, author, institution, current_date, education_level)

            # Add table of contents
            self.add_table_of_contents()

            # Add each section
            for section_name, section_content in sections.items():
                # Skip the title section as it's already added
                if section_name.lower() == "title":
                    continue

                # Add the section with appropriate formatting
                self.add_section(section_name, section_content)

            # Create output directory if it doesn't exist
            # Use a local directory for temporary storage
            output_dir = os.path.join("static/generated_papers")
            os.makedirs(output_dir, exist_ok=True)

            # Generate a unique filename
            sanitized_title = "".join(c for c in title if c.isalnum() or c.isspace()).strip()
            sanitized_title = sanitized_title.replace(" ", "_")[:50]  # Limit length
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Use paper_id in filename if available
            if paper_id:
                base_filename = f"paper_{paper_id}"
            else:
                base_filename = f"{sanitized_title}_{timestamp}"

            # Define file paths
            docx_filename = f"{base_filename}.docx"
            pdf_filename = f"{base_filename}.pdf"
            docx_filepath = os.path.join(output_dir, docx_filename)
            pdf_filepath = os.path.join(output_dir, pdf_filename)

            # Save the DOCX document
            self.document.save(docx_filepath)
            logger.info(f"Saved formatted paper (DOCX) to {docx_filepath}")

            # Try to convert to PDF using a conversion service
            pdf_filepath = self._convert_to_pdf(docx_filepath, pdf_filepath)

            # Upload to Supabase Storage if user_id and paper_id are provided
            storage_urls = {}
            if user_id and paper_id:
                try:
                    from app.services.utils.storage_service import StorageService
                    storage_service = StorageService()

                    # Upload DOCX file
                    docx_url = storage_service.upload_research_paper(
                        file_path=docx_filepath,
                        user_id=user_id,
                        paper_id=paper_id,
                        file_type="docx"
                    )

                    if docx_url:
                        storage_urls["docx"] = docx_url
                        logger.info(f"Uploaded DOCX to Supabase Storage: {docx_url}")

                    # Upload PDF file if available
                    if pdf_filepath and os.path.exists(pdf_filepath):
                        pdf_url = storage_service.upload_research_paper(
                            file_path=pdf_filepath,
                            user_id=user_id,
                            paper_id=paper_id,
                            file_type="pdf"
                        )

                        if pdf_url:
                            storage_urls["pdf"] = pdf_url
                            logger.info(f"Uploaded PDF to Supabase Storage: {pdf_url}")
                except Exception as storage_error:
                    logger.error(f"Error uploading to Supabase Storage: {str(storage_error)}")

            # Return paths to both formats
            return {
                "docx": docx_filepath,
                "pdf": pdf_filepath if os.path.exists(pdf_filepath) else None,
                "storage_urls": storage_urls
            }
        except Exception as e:
            logger.error(f"Error formatting paper: {str(e)}")
            raise

    def _convert_to_pdf(self, docx_path, pdf_path):
        """
        Convert a DOCX file to PDF.

        Args:
            docx_path: Path to the DOCX file
            pdf_path: Path to save the PDF file

        Returns:
            Path to the PDF file or None if conversion failed
        """
        try:
            # Try to use LibreOffice for conversion
            logger.info(f"Attempting to convert {docx_path} to PDF using LibreOffice")

            # Use LibreOffice headless mode for conversion
            import subprocess
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
                capture_output=True,
                text=True
            )

            if result.returncode == 0:
                logger.info(f"Successfully converted to PDF: {pdf_path}")
                return pdf_path
            else:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")

                # Try alternative conversion method using Cloudmersive API
                try:
                    from app.services.utils.cloudmersive_service import CloudmersiveService
                    cloudmersive = CloudmersiveService()

                    if cloudmersive.convert_docx_to_pdf(docx_path, pdf_path):
                        logger.info(f"Successfully converted to PDF using Cloudmersive API: {pdf_path}")
                        return pdf_path
                    else:
                        logger.error("Cloudmersive conversion failed")
                except Exception as cm_error:
                    logger.error(f"Error using Cloudmersive API: {str(cm_error)}")

                return None
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            return None
