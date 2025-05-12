"""
Document Formatter Module

This module handles the formatting of research packs into well-structured documents.
It includes support for:
- Title pages
- Table of contents with proper page numbers
- Section numbering
- Headers and footers
- Lily callout boxes
- Figures and diagrams
"""
import os
import re
import logging
import random
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configure logging
logger = logging.getLogger(__name__)

class DocumentFormatter:
    """
    Formats research pack content into a well-structured document.

    Features:
    - Professional title page
    - Table of contents with page numbers
    - Consistent section numbering
    - Headers and footers
    - Support for Lily callout boxes
    - Figure and diagram integration
    """

    def __init__(self):
        """Initialize the document formatter."""
        self.document = Document()
        self._setup_document_styles()
        self._setup_sections()

        # Section counters for proper numbering
        self.section_counters = {1: 0, 2: 0, 3: 0}

        # Track headings for TOC generation
        self.headings = []

    def _setup_document_styles(self):
        """Setup document styles for a professional academic document."""
        # Normal text style
        if 'Normal' in self.document.styles:
            normal_style = self.document.styles['Normal']
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.space_after = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5

        # Heading styles
        if 'Heading 1' in self.document.styles:
            heading1 = self.document.styles['Heading 1']
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.paragraph_format.space_before = Pt(24)
            heading1.paragraph_format.space_after = Pt(12)

        if 'Heading 2' in self.document.styles:
            heading2 = self.document.styles['Heading 2']
            heading2.font.size = Pt(14)
            heading2.font.bold = True
            heading2.paragraph_format.space_before = Pt(18)
            heading2.paragraph_format.space_after = Pt(6)

        if 'Heading 3' in self.document.styles:
            heading3 = self.document.styles['Heading 3']
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.paragraph_format.space_before = Pt(12)
            heading3.paragraph_format.space_after = Pt(6)

        # Title style
        if 'Title' in self.document.styles:
            title_style = self.document.styles['Title']
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)

        # TOC styles
        if 'TOC Heading' not in self.document.styles:
            toc_style = self.document.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_style.font.size = Pt(16)
            toc_style.font.bold = True
            toc_style.paragraph_format.space_before = Pt(24)
            toc_style.paragraph_format.space_after = Pt(12)

        if 'TOC 1' not in self.document.styles:
            toc1_style = self.document.styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
            toc1_style.font.size = Pt(12)
            toc1_style.paragraph_format.left_indent = Inches(0.25)
            toc1_style.paragraph_format.space_after = Pt(6)
            # Add tab stops for page numbers
            tab_stops = toc1_style.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # 2 is for dot leader

        if 'TOC 2' not in self.document.styles:
            toc2_style = self.document.styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
            toc2_style.font.size = Pt(11)
            toc2_style.paragraph_format.left_indent = Inches(0.5)
            toc2_style.paragraph_format.space_after = Pt(6)
            # Add tab stops for page numbers
            tab_stops = toc2_style.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # 2 is for dot leader

        # Callout style
        if 'Callout' not in self.document.styles:
            callout_style = self.document.styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
            callout_style.font.size = Pt(11)
            callout_style.paragraph_format.space_before = Pt(6)
            callout_style.paragraph_format.space_after = Pt(6)
            callout_style.paragraph_format.left_indent = Inches(0.25)
            callout_style.paragraph_format.right_indent = Inches(0.25)

    def _setup_sections(self):
        """Setup document sections for headers and footers."""
        # Ensure we have at least one section
        if not self.document.sections:
            self.document.add_section()

        # Set margins for all sections
        for section in self.document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def create_title_page(self, title: str, author: str = None, date: str = None, logo_path: str = None):
        """
        Create a professional title page.

        Args:
            title: The title of the research pack
            author: The author's name (optional)
            date: The date (optional, defaults to today)
            logo_path: Path to the logo image (optional)
        """
        # Use today's date if not provided
        if not date:
            date = datetime.now().strftime("%B %d, %Y")

        # Add logo if provided
        if logo_path and os.path.exists(logo_path):
            try:
                logo_para = self.document.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2))
                logger.info(f"Added logo from {logo_path}")
            except Exception as e:
                logger.error(f"Error adding logo: {str(e)}")

        # Add title
        title_para = self.document.add_paragraph()
        title_para.style = 'Title'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run(title).bold = True

        # Add subtitle
        subtitle_para = self.document.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.add_run("Research Pack").bold = True

        # Add spacing
        self.document.add_paragraph()

        # Add author if provided
        if author:
            author_para = self.document.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.add_run(f"Prepared for: {author}").bold = True

        # Add date
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Generated: {date}")

        # Add page break
        self.document.add_page_break()

    def add_table_of_contents(self):
        """
        Add a table of contents with page numbers.

        This creates a proper table of contents that will be updated when the document is opened.
        """
        # Add TOC heading
        toc_heading = self.document.add_heading("Table of Contents", level=1)
        toc_heading.style = 'TOC Heading'

        # Add field code for automatic TOC
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        # Add the TOC field
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # This creates TOC from level 1-3 headings with hyperlinks and page numbers

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)

        # Add page break after TOC
        self.document.add_page_break()

    def add_section(self, heading: str, content: str, level: int = 1):
        """
        Add a new section with heading and content.

        Args:
            heading: The section heading
            content: The section content
            level: The heading level (1-3)
        """
        # Add page break before main sections (level 1)
        if level == 1 and self.document.paragraphs:
            self.document.add_page_break()

        # Clean heading text
        heading = self._clean_text(heading)

        # Handle section numbering
        if level <= 3:  # Only number up to level 3
            # Special sections that shouldn't be numbered
            special_sections = ["references", "abstract", "appendix", "acknowledgements", "table of contents"]

            if heading.lower() not in special_sections:
                # Update section counters
                if level == 1:
                    self.section_counters[1] += 1
                    self.section_counters[2] = 0
                    self.section_counters[3] = 0
                    section_number = f"{self.section_counters[1]}"
                elif level == 2:
                    self.section_counters[2] += 1
                    self.section_counters[3] = 0
                    section_number = f"{self.section_counters[1]}.{self.section_counters[2]}"
                elif level == 3:
                    self.section_counters[3] += 1
                    section_number = f"{self.section_counters[1]}.{self.section_counters[2]}.{self.section_counters[3]}"

                # Add numbered heading
                heading_with_number = f"{section_number}. {heading}"
            else:
                heading_with_number = heading
        else:
            heading_with_number = heading

        # Add the heading
        heading_para = self.document.add_heading(heading_with_number, level=level)

        # Track headings for TOC
        self.headings.append((heading_with_number, level))

        # Process content for Lily callouts
        if content and "[[LILY_CALLOUT" in content:
            self._process_content_with_callouts(content)
        else:
            # Add regular content
            if content:
                content = self._clean_text(content)
                content_para = self.document.add_paragraph(content)
                content_para.style = 'Normal'

    def _process_content_with_callouts(self, content: str):
        """
        Process content that contains Lily callout markers.

        Args:
            content: The content with Lily callout markers
        """
        # Define the pattern for Lily callouts
        pattern = r'\[\[LILY_CALLOUT\s+type="(\w+)"\s+title="([^"]+)"(?:\s+source="([^"]+)")?\s*\]\]([\s\S]*?)\[\[\/LILY_CALLOUT\]\]'

        # Find all callouts
        matches = list(re.finditer(pattern, content, re.DOTALL))

        if not matches:
            # No callouts found, add as regular paragraph
            content_para = self.document.add_paragraph(self._clean_text(content))
            content_para.style = 'Normal'
            return

        # Process content with callouts
        last_end = 0

        for match in matches:
            # Add text before this callout
            if match.start() > last_end:
                before_text = content[last_end:match.start()].strip()
                if before_text:
                    before_para = self.document.add_paragraph(self._clean_text(before_text))
                    before_para.style = 'Normal'

            # Extract callout information
            callout_type = match.group(1)
            callout_title = match.group(2)
            callout_source = match.group(3)
            callout_content = match.group(4).strip()

            # Add the callout
            self.add_lily_callout(callout_title, callout_content, callout_type)

            # Update last_end
            last_end = match.end()

        # Add text after the last callout
        if last_end < len(content):
            after_text = content[last_end:].strip()
            if after_text:
                after_para = self.document.add_paragraph(self._clean_text(after_text))
                after_para.style = 'Normal'

    def add_lily_callout(self, title: str, content: str, callout_type: str = "insight"):
        """
        Add a Lily callout box.

        Args:
            title: The callout title
            content: The callout content
            callout_type: The type of callout (tip, insight, question, warning, etc.)
        """
        # Create a table for the callout
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Light Grid'

        # Get the cell
        cell = table.cell(0, 0)

        # Add title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"Lily's {callout_type.capitalize()}: {title}")
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Add content
        content_para = cell.add_paragraph()
        content_para.text = self._clean_text(content)
        content_para.style = 'Normal'

        # Add spacing after callout
        self.document.add_paragraph()

        return table

    def add_figure(self, image_path: str, caption: str = None):
        """
        Add a figure with caption.

        Args:
            image_path: Path to the image file
            caption: Optional caption for the figure
        """
        try:
            # Check if image exists
            if not os.path.exists(image_path):
                logger.error(f"Image not found: {image_path}")
                return

            # Add the image
            figure_para = self.document.add_paragraph()
            figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure_run = figure_para.add_run()
            figure_run.add_picture(image_path, width=Inches(6))

            # Add caption if provided
            if caption:
                caption_para = self.document.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.style = 'Caption'
                caption_para.add_run(f"Figure: {caption}").italic = True

            logger.info(f"Added figure from {image_path}")

        except Exception as e:
            logger.error(f"Error adding figure: {str(e)}")

    def add_header_footer(self, title: str):
        """
        Add header and footer to the document.

        Args:
            title: The document title for the header
        """
        # Get a random quote
        quote = self._get_random_quote()

        # Get footer text from config
        from app.services.config.config_service import ConfigService
        config_service = ConfigService()
        footer_text = config_service.get_config('footer_copyright_text',
                                           "© " + str(datetime.now().year) + " researchassistant.uk - All rights reserved")

        # Add header and footer to each section
        for section in self.document.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = title
            header_para.style = 'Header'
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

            # Add quote
            quote_run = footer_para.add_run(f"{quote}   |   ")
            quote_run.font.size = Pt(9)
            quote_run.italic = True

            # Add copyright from config
            copyright_run = footer_para.add_run(footer_text + " | ")
            copyright_run.font.size = Pt(9)

            # Add page number
            page_run = footer_para.add_run("Page ")
            page_run.font.size = Pt(9)
            self._add_page_number(page_run)

            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_page_number(self, run):
        """
        Add a page number field to a run.

        Args:
            run: The run to add the page number to
        """
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)

    def add_paragraph(self, text: str):
        """
        Add a paragraph to the document.

        Args:
            text: The text to add
        """
        if not text:
            return

        # Clean the text
        text = self._clean_text(text)

        # Add the paragraph
        para = self.document.add_paragraph(text)
        para.style = 'Normal'

        return para

    def _get_random_quote(self) -> str:
        """
        Get a random Lily quote from the database.

        Returns:
            A random quote or a default quote if none are found
        """
        try:
            # Import the config service
            from app.services.config import config_service

            if not hasattr(config_service, 'db_connection') or not config_service.db_connection:
                logger.warning("No database connection available. Using default quote.")
                return "Keep pushing forward—breakthroughs come from persistence!"

            # Query the quotes table
            query = """
            SELECT quote
            FROM saas_lily_quotes
            WHERE enabled = true
            ORDER BY RANDOM()
            LIMIT 1
            """

            result = config_service.db_connection.execute(query).fetchone()

            if result:
                return result[0]
            else:
                logger.warning("No quotes found in database. Using default quote.")
                return "Keep pushing forward—breakthroughs come from persistence!"

        except Exception as e:
            logger.error(f"Error retrieving random quote: {str(e)}")
            return "Keep pushing forward—breakthroughs come from persistence!"

    def _clean_text(self, text: str) -> str:
        """
        Clean text by removing markdown formatting.

        Args:
            text: The text to clean

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove markdown formatting
        text = re.sub(r'#+ ', '', text)  # Remove heading markers
        text = re.sub(r'\*\*|\*|__|_', '', text)  # Remove bold/italic markers
        text = re.sub(r'```|`', '', text)  # Remove code markers

        return text.strip()

    def format_research_pack(self, research_pack: Dict) -> Document:
        """
        Format a complete research pack into a well-structured document.

        Args:
            research_pack: The research pack data

        Returns:
            The formatted document
        """
        # Reset document
        self.document = Document()
        self._setup_document_styles()
        self._setup_sections()
        self.section_counters = {1: 0, 2: 0, 3: 0}
        self.headings = []

        # Extract key information
        title = research_pack.get("title", "Research Pack")
        author = research_pack.get("author", "")
        date = research_pack.get("date_generated", datetime.now().strftime("%B %d, %Y"))

        # Get the logo path
        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "images", "coversheet-lilylogo.png")

        # Check if the logo exists
        if not os.path.exists(logo_path):
            logger.warning(f"Logo file not found at {logo_path}, trying alternative paths")
            # Try alternative paths
            alt_paths = [
                os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "images", "logo-lily.png"),
                os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "logo-lily.png"),
                os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "assets", "images", "logo-lily.png")
            ]

            for alt_path in alt_paths:
                if os.path.exists(alt_path):
                    logo_path = alt_path
                    logger.info(f"Found logo at alternative path: {logo_path}")
                    break
            else:
                logger.error("Logo file not found in any of the expected locations")
                logo_path = None

        # Create title page with logo
        self.create_title_page(title, author, date, logo_path)

        # Add table of contents
        self.add_table_of_contents()

        # Add header and footer
        self.add_header_footer(title)

        # Process sections
        sections = research_pack.get("sections", {})
        diagrams = research_pack.get("diagrams", [])

        # Get section configuration from database
        section_config = self._get_section_configuration()
        logger.info(f"Retrieved {len(section_config)} section configurations from database")

        # Process sections based on database configuration
        self._process_sections_from_config(sections, diagrams, section_config)

        # Section 1: Introduction (500-750 words)
        if "introduction" in sections:
            self.add_section("Introduction", "", 1)

            # Add introduction content
            intro_content = sections["introduction"]
            if isinstance(intro_content, dict):
                # Check if there's a diagram in the section
                section_diagram = intro_content.get("diagram")

                # Process other content
                for key, content in intro_content.items():
                    if key != "diagram" and key != "comparative_diagram":
                        title = " ".join(word.capitalize() for word in key.split("_"))
                        self.add_section(title, content, 2)

                # Add the section diagram if available
                if section_diagram and isinstance(section_diagram, dict):
                    image_path = section_diagram.get("path", "")
                    caption = section_diagram.get("title", "Question Breakdown")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding question breakdown diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Question breakdown diagram file not found at {image_path}")
            else:
                self.add_paragraph(intro_content)

            # Add question breakdown diagram if available from diagrams list (fallback)
            if not (isinstance(intro_content, dict) and "diagram" in intro_content):
                question_breakdown = next((d for d in diagrams if d.get("type") == "question_breakdown"), None)
                if question_breakdown:
                    image_path = question_breakdown.get("path", "")
                    caption = question_breakdown.get("caption", question_breakdown.get("title", "Question Breakdown"))

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding question breakdown diagram from diagrams list: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Question breakdown diagram file not found at {image_path}")

        # Section 2: Topic Analysis (5000-7000 words)
        if "topic_analysis" in sections:
            self.add_section("Topic Analysis", "", 1)

            # Add topic analysis content
            topic_content = sections["topic_analysis"]
            if isinstance(topic_content, dict):
                # Check if there's a diagram in the section
                section_diagram = topic_content.get("diagram")

                # Process other content
                for key, content in topic_content.items():
                    if key != "diagram" and key != "comparative_diagram":
                        title = " ".join(word.capitalize() for word in key.split("_"))
                        self.add_section(title, content, 2)

                # Add the section diagram if available
                if section_diagram and isinstance(section_diagram, dict):
                    image_path = section_diagram.get("path", "")
                    caption = section_diagram.get("title", "Topic Analysis Mind Map")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding section diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Section diagram file not found at {image_path}")
            else:
                self.add_paragraph(topic_content)

            # Add mind map diagram if available from diagrams list (fallback)
            if not (isinstance(topic_content, dict) and "diagram" in topic_content):
                mind_map_diagram = next((d for d in diagrams if d.get("type") == "mind_map"), None)
                if mind_map_diagram:
                    image_path = mind_map_diagram.get("path", "")
                    caption = mind_map_diagram.get("caption", mind_map_diagram.get("title", "Topic Analysis Mind Map"))

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding mind map diagram from diagrams list: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Mind map diagram file not found at {image_path}")

        # Section 3: Methodological Approaches (3000-4000 words)
        if "methodological_approaches" in sections:
            self.add_section("Methodological Approaches", "", 1)

            # Add methodological approaches content
            method_content = sections["methodological_approaches"]
            if isinstance(method_content, dict):
                # Check if there's a diagram in the section
                section_diagram = method_content.get("diagram")

                # Process other content
                for key, content in method_content.items():
                    if key != "diagram" and key != "comparative_diagram":
                        title = " ".join(word.capitalize() for word in key.split("_"))
                        self.add_section(title, content, 2)

                # Add the section diagram if available
                if section_diagram and isinstance(section_diagram, dict):
                    image_path = section_diagram.get("path", "")
                    caption = section_diagram.get("title", "Research Methodology Process Flow")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding section diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Section diagram file not found at {image_path}")
            else:
                self.add_paragraph(method_content)

            # Add process flow or journey map diagram if available from diagrams list (fallback)
            if not (isinstance(method_content, dict) and "diagram" in method_content):
                # Try process flow first
                process_diagram = next((d for d in diagrams if d.get("type") == "process_flow"), None)
                if process_diagram:
                    image_path = process_diagram.get("path", "")
                    caption = process_diagram.get("caption", process_diagram.get("title", "Research Methodology Process Flow"))

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding process flow diagram from diagrams list: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Process flow diagram file not found at {image_path}")
                else:
                    # Try journey map as an alternative
                    journey_map = next((d for d in diagrams if d.get("type") == "journey_map"), None)
                    if journey_map:
                        image_path = journey_map.get("path", "")
                        caption = journey_map.get("caption", journey_map.get("title", "Research Journey Map"))

                        if image_path and os.path.exists(image_path):
                            logger.info(f"Adding journey map diagram from diagrams list: {image_path}")
                            self.add_figure(image_path, caption)
                        else:
                            logger.warning(f"Journey map diagram file not found at {image_path}")

        # Section 4: Key Arguments (5000-7000 words)
        if "key_arguments" in sections:
            self.add_section("Key Arguments", "", 1)

            # Add key arguments content
            arguments_content = sections["key_arguments"]
            if isinstance(arguments_content, dict):
                # Check if there are diagrams in the section
                section_diagram = arguments_content.get("diagram")
                comparative_diagram = arguments_content.get("comparative_diagram")

                # Process other content
                for key, content in arguments_content.items():
                    if key != "diagram" and key != "comparative_diagram":
                        title = " ".join(word.capitalize() for word in key.split("_"))
                        self.add_section(title, content, 2)

                # Add the main section diagram if available
                if section_diagram and isinstance(section_diagram, dict):
                    image_path = section_diagram.get("path", "")
                    caption = section_diagram.get("title", "Argument Mapping")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding argument mapping diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Argument mapping diagram file not found at {image_path}")

                # Add the comparative diagram if available
                if comparative_diagram and isinstance(comparative_diagram, dict):
                    image_path = comparative_diagram.get("path", "")
                    caption = comparative_diagram.get("title", "Comparative Analysis of Key Arguments")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding comparative analysis diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Comparative analysis diagram file not found at {image_path}")
            else:
                self.add_paragraph(arguments_content)

            # Add diagrams from diagrams list if not already added (fallback)
            has_argument_mapping = isinstance(arguments_content, dict) and "diagram" in arguments_content
            has_comparative = isinstance(arguments_content, dict) and "comparative_diagram" in arguments_content

            # Add argument mapping if not already added
            if not has_argument_mapping:
                argument_mapping = next((d for d in diagrams if d.get("type") == "argument_mapping"), None)
                if argument_mapping:
                    image_path = argument_mapping.get("path", "")
                    caption = argument_mapping.get("caption", argument_mapping.get("title", "Argument Mapping"))

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding argument mapping diagram from diagrams list: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Argument mapping diagram file not found at {image_path}")

            # Add comparative analysis if not already added
            if not has_comparative:
                comparative_diagram = next((d for d in diagrams if d.get("type") == "comparative_analysis"), None)
                if comparative_diagram:
                    image_path = comparative_diagram.get("path", "")
                    caption = comparative_diagram.get("caption", comparative_diagram.get("title", "Comparative Analysis of Key Arguments"))

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding comparative analysis diagram from diagrams list: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Comparative analysis diagram file not found at {image_path}")

        # Section 5: Citations and Resources (1000-2000 words)
        if "citations_resources" in sections:
            self.add_section("Citations and Resources", "", 1)

            # Add citations and resources content
            citations_content = sections["citations_resources"]
            if isinstance(citations_content, dict):
                for key, content in citations_content.items():
                    title = " ".join(word.capitalize() for word in key.split("_"))
                    self.add_section(title, content, 2)
            else:
                self.add_paragraph(citations_content)

        # Handle legacy section names for backward compatibility
        # This ensures that research packs created with the old structure still format correctly

        # Legacy: Topic Analysis (if not already processed)
        if "topic_analysis" not in sections and "topic_analysis_legacy" in sections:
            self.add_section("Topic Analysis", "", 1)
            topic_content = sections["topic_analysis_legacy"]
            if isinstance(topic_content, dict):
                for key, content in topic_content.items():
                    title = " ".join(word.capitalize() for word in key.split("_"))
                    self.add_section(title, content, 2)
            else:
                self.add_paragraph(topic_content)

        # Legacy: Methodology (if not already processed)
        if "methodological_approaches" not in sections and "methodology" in sections:
            self.add_section("Methodological Approaches", "", 1)
            method_content = sections["methodology"]
            if isinstance(method_content, dict):
                for key, content in method_content.items():
                    title = " ".join(word.capitalize() for word in key.split("_"))
                    self.add_section(title, content, 2)
            else:
                self.add_paragraph(method_content)

        # Legacy: Citations (if not already processed)
        if "citations_resources" not in sections and "citations" in sections:
            self.add_section("Citations and Resources", "", 1)
            citations_content = sections["citations"]
            if isinstance(citations_content, dict):
                for key, content in citations_content.items():
                    title = " ".join(word.capitalize() for word in key.split("_"))
                    self.add_section(title, content, 2)
            else:
                self.add_paragraph(citations_content)

        # Add any remaining diagrams that weren't included in specific sections
        remaining_diagrams = [d for d in diagrams if
                             d.get("type") not in ["mind_map", "process_flow", "comparative_analysis"]]

        # Log all diagrams for debugging
        logger.info(f"Document formatter: Processing {len(diagrams)} diagrams")
        for i, diagram in enumerate(diagrams):
            diagram_type = diagram.get("type", "unknown")
            diagram_path = diagram.get("path", "unknown")
            logger.info(f"Document formatter: Diagram {i+1}: type={diagram_type}, path={diagram_path}")

            # Check if the diagram file exists
            if diagram_path and os.path.exists(diagram_path):
                logger.info(f"Document formatter: Diagram file exists at {diagram_path}")
            else:
                logger.warning(f"Document formatter: Diagram file does not exist at {diagram_path}")

        # Add diagrams section if we have diagrams but they weren't included in specific sections
        if diagrams and not any(d.get("type") in ["mind_map", "process_flow", "comparative_analysis"] for d in diagrams):
            self.add_section("Research Diagrams", "", 1)

            # Add each diagram
            for i, diagram in enumerate(diagrams):
                image_path = diagram.get("path", "")
                caption = diagram.get("caption", f"Diagram {i+1}")
                diagram_type = diagram.get("type", "general")

                if image_path and os.path.exists(image_path):
                    logger.info(f"Document formatter: Adding diagram {i+1} to document")
                    self.add_figure(image_path, f"{caption} ({diagram_type.replace('_', ' ').title()})")
                else:
                    logger.warning(f"Document formatter: Cannot add diagram {i+1}, file not found at {image_path}")

        # Appendices section
        if "appendices" in sections:
            self.add_section("Appendices", "", 1)

            # Add appendices content
            appendices_content = sections["appendices"].get("content", "")
            if appendices_content:
                self.add_section("Research Tools and Templates", appendices_content, 2)

        if remaining_diagrams:
            self.add_section("Additional Diagrams", "", 1)

            for i, diagram in enumerate(remaining_diagrams):
                image_path = diagram.get("path", "")
                caption = diagram.get("caption", f"Diagram {i+1}")
                diagram_type = diagram.get("type", "general")

                if image_path and os.path.exists(image_path):
                    self.add_figure(image_path, f"{caption} ({diagram_type.replace('_', ' ').title()})")

        return self.document

    def add_paragraph(self, text: str = ""):
        """
        Add a paragraph to the document.

        Args:
            text: The paragraph text
        """
        para = self.document.add_paragraph(self._clean_text(text))
        para.style = 'Normal'
        return para

    def save_document(self, filepath: str) -> str:
        """
        Save the document to a file.

        Args:
            filepath: The path to save the document to

        Returns:
            The actual filepath used
        """
        try:
            # Ensure directory exists
            directory = os.path.dirname(filepath)
            if directory and not os.path.exists(directory):
                os.makedirs(directory)

            # Save the document
            self.document.save(filepath)
            logger.info(f"Document saved to {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")

            # Try to save with a different filename
            try:
                # Generate a random filename
                random_suffix = ''.join(random.choices('abcdefghijklmnopqrstuvwxyz0123456789', k=8))
                new_filepath = f"{os.path.splitext(filepath)[0]}_{random_suffix}.docx"

                self.document.save(new_filepath)
                logger.info(f"Document saved to alternate path: {new_filepath}")
                return new_filepath

            except Exception as e2:
                logger.error(f"Error saving document to alternate path: {str(e2)}")
                raise

    def format_paper(self, paper: Dict, author: str, institution: str, education_level: str, user_id: str = None, paper_id: str = None) -> Dict:
        """
        Format a paper into a professional document.

        Args:
            paper: The paper object containing title, sections, etc.
            author: The author's name
            institution: The institution name
            education_level: The education level (university, postgraduate, etc.)
            user_id: The ID of the user who owns the paper
            paper_id: The ID of the paper

        Returns:
            Dictionary with paths to the saved documents (docx and pdf)
        """
        try:
            # Reset document
            self.document = Document()
            self._setup_document_styles()
            self._setup_sections()
            self.section_counters = {1: 0, 2: 0, 3: 0}
            self.headings = []

            # Get the title and sections from the paper
            title = paper.get("title", "Research Paper")
            sections = paper.get("sections", {})

            # Create the title page
            current_date = datetime.now().strftime("%B %d, %Y")

            # Get the logo path
            logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "images", "coversheet-lilylogo.png")

            # Check if the logo exists
            if not os.path.exists(logo_path):
                logger.warning(f"Logo file not found at {logo_path}, trying alternative paths")
                # Try alternative paths
                alt_paths = [
                    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "images", "logo-lily.png"),
                    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "static", "logo-lily.png"),
                    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "assets", "images", "logo-lily.png")
                ]

                for alt_path in alt_paths:
                    if os.path.exists(alt_path):
                        logo_path = alt_path
                        logger.info(f"Found logo at alternative path: {logo_path}")
                        break
                else:
                    logger.error("Logo file not found in any of the expected locations")
                    logo_path = None

            # Create title page with logo
            self.create_title_page(title, author, institution, current_date, logo_path)

            # Add table of contents
            self.add_table_of_contents()

            # Add header and footer
            self.add_header_footer(title)

            # Add each section
            for section_name, section_content in sections.items():
                # Skip the title section as it's already added
                if section_name.lower() == "title":
                    continue

                # Add the section with appropriate formatting
                self.add_section(section_name, section_content, 1)

            # Create output directory if it doesn't exist
            # Use a local directory for temporary storage
            output_dir = os.path.join("static/generated_papers")
            os.makedirs(output_dir, exist_ok=True)

            # Generate a unique filename
            sanitized_title = "".join(c for c in title if c.isalnum() or c.isspace()).strip()
            sanitized_title = sanitized_title.replace(" ", "_")[:50]  # Limit length
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Use paper_id in filename if available
            if paper_id:
                base_filename = f"paper_{paper_id}"
            else:
                base_filename = f"{sanitized_title}_{timestamp}"

            # Define file paths
            docx_filename = f"{base_filename}.docx"
            pdf_filename = f"{base_filename}.pdf"
            docx_filepath = os.path.join(output_dir, docx_filename)
            pdf_filepath = os.path.join(output_dir, pdf_filename)

            # Save the DOCX document
            self.document.save(docx_filepath)
            logger.info(f"Saved formatted paper (DOCX) to {docx_filepath}")

            # Try to convert to PDF using a conversion service
            pdf_filepath = self._convert_to_pdf(docx_filepath, pdf_filepath)

            # Upload to Supabase Storage if user_id and paper_id are provided
            storage_urls = {}
            if user_id and paper_id:
                try:
                    from app.services.utils.storage_service import StorageService
                    storage_service = StorageService()

                    # Upload DOCX file
                    docx_url = storage_service.upload_research_paper(
                        file_path=docx_filepath,
                        user_id=user_id,
                        paper_id=paper_id,
                        file_type="docx"
                    )

                    if docx_url:
                        storage_urls["docx"] = docx_url
                        logger.info(f"Uploaded DOCX to Supabase Storage: {docx_url}")

                    # Upload PDF file if available
                    if pdf_filepath and os.path.exists(pdf_filepath):
                        pdf_url = storage_service.upload_research_paper(
                            file_path=pdf_filepath,
                            user_id=user_id,
                            paper_id=paper_id,
                            file_type="pdf"
                        )

                        if pdf_url:
                            storage_urls["pdf"] = pdf_url
                            logger.info(f"Uploaded PDF to Supabase Storage: {pdf_url}")
                except Exception as storage_error:
                    logger.error(f"Error uploading to Supabase Storage: {str(storage_error)}")

            # Return paths to both formats
            return {
                "docx": docx_filepath,
                "pdf": pdf_filepath if os.path.exists(pdf_filepath) else None,
                "storage_urls": storage_urls
            }
        except Exception as e:
            logger.error(f"Error formatting paper: {str(e)}")
            raise

    def _convert_to_pdf(self, docx_path, pdf_path):
        """
        Convert a DOCX file to PDF.

        Args:
            docx_path: Path to the DOCX file
            pdf_path: Path to save the PDF file

        Returns:
            Path to the PDF file or None if conversion failed
        """
        try:
            # Try to use LibreOffice for conversion
            logger.info(f"Attempting to convert {docx_path} to PDF using LibreOffice")

            # Use LibreOffice headless mode for conversion
            import subprocess
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
                capture_output=True,
                text=True
            )

            if result.returncode == 0:
                logger.info(f"Successfully converted to PDF: {pdf_path}")
                return pdf_path
            else:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")

                # Try alternative conversion method using Cloudmersive API
                try:
                    from app.services.utils.cloudmersive_service import CloudmersiveService
                    cloudmersive = CloudmersiveService()

                    if cloudmersive.convert_docx_to_pdf(docx_path, pdf_path):
                        logger.info(f"Successfully converted to PDF using Cloudmersive API: {pdf_path}")
                        return pdf_path
                    else:
                        logger.error("Cloudmersive conversion failed")
                except Exception as cm_error:
                    logger.error(f"Error using Cloudmersive API: {str(cm_error)}")

                return None
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            return None

    def _get_section_configuration(self) -> List[Dict]:
        """
        Get section configuration from the database.

        Returns:
            A list of section configurations
        """
        try:
            # Import here to avoid circular imports
            from app.utils.supabase_client import get_supabase_client

            # Get Supabase client
            supabase_client = get_supabase_client()
            if not supabase_client:
                logger.error("Failed to get Supabase client")
                return []

            # Query the database for section configuration
            response = supabase_client.table('saas_researchpack_sections').select('*').order('section_order').execute()

            if response.data:
                logger.info(f"Retrieved {len(response.data)} section configurations from database")
                return response.data
            else:
                logger.warning("No section configurations found in database")
                return []
        except Exception as e:
            logger.error(f"Error retrieving section configuration from database: {str(e)}")
            return []

    def _process_sections_from_config(self, sections: Dict, diagrams: List[Dict], section_config: List[Dict]) -> None:
        """
        Process sections based on configuration from the database.

        Args:
            sections: The sections data
            diagrams: The diagrams data
            section_config: The section configuration from the database
        """
        # If no section configuration, use default processing
        if not section_config:
            logger.warning("No section configuration found, using default processing")
            self._process_sections_default(sections, diagrams)
            return

        # Process sections based on configuration
        for config in section_config:
            section_name = config.get('name')  # Using 'name' instead of 'section_name'
            display_name = config.get('display_name')
            include_diagrams = config.get('include_diagrams', False)
            diagram_type = config.get('diagram_type')

            # Skip if section is not enabled
            if not config.get('enabled', True):
                logger.info(f"Skipping disabled section: {section_name}")
                continue

            # Skip if section is not in the content
            if section_name not in sections:
                logger.info(f"Section {section_name} not found in content")
                continue

            # Add the section heading
            self.add_section(display_name, "", 1)
            logger.info(f"Added section: {display_name}")

            # Process section content
            section_content = sections[section_name]
            if isinstance(section_content, dict):
                # Check for diagrams in the section
                section_diagram = section_content.get("diagram")
                comparative_diagram = section_content.get("comparative_diagram")

                # Process subsections
                for key, content in section_content.items():
                    if key != "diagram" and key != "comparative_diagram":
                        title = " ".join(word.capitalize() for word in key.split("_"))
                        self.add_section(title, content, 2)
                        logger.info(f"Added subsection: {title}")

                # Add section diagram if available
                if section_diagram and isinstance(section_diagram, dict):
                    image_path = section_diagram.get("path", "")
                    caption = section_diagram.get("title", f"{display_name} Diagram")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding section diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Section diagram file not found at {image_path}")

                # Add comparative diagram if available
                if comparative_diagram and isinstance(comparative_diagram, dict):
                    image_path = comparative_diagram.get("path", "")
                    caption = comparative_diagram.get("title", f"Comparative Analysis for {display_name}")

                    if image_path and os.path.exists(image_path):
                        logger.info(f"Adding comparative diagram from path: {image_path}")
                        self.add_figure(image_path, caption)
                    else:
                        logger.warning(f"Comparative diagram file not found at {image_path}")
            elif isinstance(section_content, str):
                self.add_paragraph(section_content)
                logger.info(f"Added content for section: {display_name}")
            elif isinstance(section_content, dict) and "content" in section_content:
                self.add_paragraph(section_content["content"])
                logger.info(f"Added content for section: {display_name}")

            # Add diagram from diagrams list if configured and not already added
            if include_diagrams and diagram_type:
                has_diagram = (isinstance(section_content, dict) and
                              ("diagram" in section_content or
                               (diagram_type == "comparative_analysis" and "comparative_diagram" in section_content)))

                if not has_diagram:
                    matching_diagram = next((d for d in diagrams if d.get("type") == diagram_type), None)
                    if matching_diagram:
                        image_path = matching_diagram.get("path", "")
                        caption = matching_diagram.get("caption", matching_diagram.get("title", f"{display_name} {diagram_type.replace('_', ' ').title()}"))

                        if image_path and os.path.exists(image_path):
                            logger.info(f"Adding {diagram_type} diagram from diagrams list: {image_path}")
                            self.add_figure(image_path, caption)
                        else:
                            logger.warning(f"{diagram_type} diagram file not found at {image_path}")

        # Add any remaining diagrams that weren't included in specific sections
        remaining_diagrams = [d for d in diagrams if not any(config.get('diagram_type') == d.get("type") for config in section_config if config.get('include_diagrams', False))]

        if remaining_diagrams:
            self.add_section("Additional Diagrams", "", 1)
            logger.info(f"Adding {len(remaining_diagrams)} remaining diagrams")

            for i, diagram in enumerate(remaining_diagrams):
                image_path = diagram.get("path", "")
                caption = diagram.get("caption", diagram.get("title", f"Diagram {i+1}"))
                diagram_type = diagram.get("type", "general")

                if image_path and os.path.exists(image_path):
                    self.add_figure(image_path, f"{caption} ({diagram_type.replace('_', ' ').title()})")
                    logger.info(f"Added remaining diagram: {diagram_type}")
                else:
                    logger.warning(f"Remaining diagram file not found at {image_path}")

    def _process_sections_default(self, sections: Dict, diagrams: List[Dict]) -> None:
        """
        Process sections using the default approach (fallback if no database configuration).

        Args:
            sections: The sections data
            diagrams: The diagrams data
        """
        logger.warning("Using default section processing (not database-driven)")

        # Static sections
        static_sections = {
            "about_lily": "About Lily AI",
            "how_to_use": "How to Use This Pack",
            "appendices": "Appendices"
        }

        # Process static sections
        for section_name, display_name in static_sections.items():
            if section_name in sections:
                content = sections[section_name].get("content", "")
                self.add_section(display_name, content, 1)
                logger.info(f"Added static section: {display_name}")

        # Process dynamic sections with standard structure
        standard_sections = {
            "introduction": "Introduction",
            "topic_analysis": "Topic Analysis",
            "methodological_approaches": "Methodological Approaches",
            "key_arguments": "Key Arguments",
            "citations_resources": "Citations and Resources",
            "personalized_questions": "Personalized Questions"
        }

        # Map sections to diagram types
        section_diagram_types = {
            "introduction": "question_breakdown",
            "topic_analysis": "mind_map",
            "methodological_approaches": ["process_flow", "journey_map"],
            "key_arguments": ["argument_mapping", "comparative_analysis"]
        }

        # Process standard sections
        for section_name, display_name in standard_sections.items():
            if section_name in sections:
                # Add section heading
                self.add_section(display_name, "", 1)
                logger.info(f"Added standard section: {display_name}")

                # Process section content
                section_content = sections[section_name]
                if isinstance(section_content, dict):
                    # Handle personalized questions specially
                    if section_name == "personalized_questions":
                        questions_data = section_content.get("questions", {})
                        if questions_data:
                            for key, data in questions_data.items():
                                question = data.get("question", "")
                                answer = data.get("answer", "")

                                # Add the question as a subheading
                                self.add_section(f"Q: {question}", "", 2)

                                # Add the answer
                                self.add_paragraph(answer)

                                # Add some space after each Q&A
                                self.document.add_paragraph()
                    else:
                        # Process subsections
                        for key, content in section_content.items():
                            if key != "diagram" and key != "comparative_diagram":
                                title = " ".join(word.capitalize() for word in key.split("_"))
                                self.add_section(title, content, 2)
                                logger.info(f"Added subsection: {title}")
                else:
                    self.add_paragraph(section_content)

                # Add diagrams if applicable
                diagram_types = section_diagram_types.get(section_name)
                if diagram_types:
                    if not isinstance(diagram_types, list):
                        diagram_types = [diagram_types]

                    for diagram_type in diagram_types:
                        matching_diagram = next((d for d in diagrams if d.get("type") == diagram_type), None)
                        if matching_diagram:
                            image_path = matching_diagram.get("path", "")
                            caption = matching_diagram.get("caption", matching_diagram.get("title", f"{display_name} {diagram_type.replace('_', ' ').title()}"))

                            if image_path and os.path.exists(image_path):
                                logger.info(f"Adding {diagram_type} diagram: {image_path}")
                                self.add_figure(image_path, caption)
                            else:
                                logger.warning(f"{diagram_type} diagram file not found at {image_path}")

        # Add any remaining diagrams
        remaining_diagrams = [d for d in diagrams if
                             d.get("type") not in ["mind_map", "process_flow", "journey_map",
                                                  "argument_mapping", "comparative_analysis",
                                                  "question_breakdown"]]

        if remaining_diagrams:
            self.add_section("Additional Diagrams", "", 1)

            for i, diagram in enumerate(remaining_diagrams):
                image_path = diagram.get("path", "")
                caption = diagram.get("caption", f"Diagram {i+1}")
                diagram_type = diagram.get("type", "general")

                if image_path and os.path.exists(image_path):
                    self.add_figure(image_path, f"{caption} ({diagram_type.replace('_', ' ').title()})")
