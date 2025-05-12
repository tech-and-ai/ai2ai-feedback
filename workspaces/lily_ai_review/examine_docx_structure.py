import docx
import sys

def examine_docx_structure(file_path):
    doc = docx.Document(file_path)
    
    print(f"Document has {len(doc.paragraphs)} paragraphs")
    print(f"Document has {len(doc.tables)} tables")
    
    # Examine paragraph styles and formatting
    print("\n=== PARAGRAPH STYLES AND FORMATTING ===\n")
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style_name = para.style.name if para.style else "No style"
            print(f"Paragraph {i}: Style: {style_name} | Text: {para.text[:50]}...")
            
            # Check for runs with different formatting
            if len(para.runs) > 1:
                print(f"  Paragraph {i} has {len(para.runs)} runs with different formatting:")
                for j, run in enumerate(para.runs):
                    if run.text.strip():
                        bold = "Bold" if run.bold else "Not bold"
                        italic = "Italic" if run.italic else "Not italic"
                        underline = "Underlined" if run.underline else "Not underlined"
                        print(f"    Run {j}: {bold}, {italic}, {underline} | Text: {run.text[:30]}...")
    
    # Examine table styles
    if doc.tables:
        print("\n=== TABLE STYLES ===\n")
        for i, table in enumerate(doc.tables):
            style_name = table.style.name if hasattr(table, 'style') and table.style else "No style"
            print(f"Table {i}: Style: {style_name}")
            # Print first row as sample
            if table.rows:
                row_text = [cell.text for cell in table.rows[0].cells]
                print(f"  First row: {' | '.join(row_text)}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python examine_docx_structure.py <path_to_docx_file>")
        sys.exit(1)
    
    examine_docx_structure(sys.argv[1])
